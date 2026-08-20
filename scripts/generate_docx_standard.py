import os
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

BASE_DIR = r"D:\Tawanagent\TAWAN-OS\02_STUDY\2026-Semester\Cloud_Computing\Cloud_Project\DISNEY_LORCANA_PLAYLAB_CLOUD"
DOCS_DIR = os.path.join(BASE_DIR, "docs", "reports")
IMG_DIR = os.path.join(DOCS_DIR, "images")
OUT_DOCX = os.path.join(DOCS_DIR, "G21_รูปเล่มรายงาน_DISNEY_LORCANA_CLOUD.docx")

doc = docx.Document()

# Set standard margins (2.54 cm = 1 inch)
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header.is_linked_to_previous = False
    p_hdr = section.header.paragraphs[0]
    p_hdr.text = "กลุ่ม G21: Disney Lorcana PlayLab Cloud | รายงานความก้าวหน้า Stage 2 (Sprint 1–3)"
    p_hdr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if p_hdr.runs:
        p_hdr.runs[0].font.name = "TH Sarabun New"
        p_hdr.runs[0].font.size = Pt(10)
        p_hdr.runs[0].font.italic = True
        p_hdr.runs[0].font.color.rgb = RGBColor(120, 120, 120)

# Helper function to set run font style
def set_font(run, font_name="TH Sarabun New", size_pt=12, bold=False, italic=False, color_rgb=None):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    if color_rgb:
        run.font.color.rgb = color_rgb
    
    # Force eastAsia & ascii / hAnsi font in XML
    rPr = run._r.get_or_add_rPr()
    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{font_name}" w:hAnsi="{font_name}" w:cs="{font_name}" w:eastAsia="{font_name}"/>')
    rPr.append(rFonts)

def add_heading_1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_font(run, size_pt=15, bold=True, color_rgb=RGBColor(24, 43, 73)) # 20px = 15pt
    return p

def add_heading_2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_font(run, size_pt=13.5, bold=True, color_rgb=RGBColor(40, 70, 110)) # 18px = 13.5pt
    return p

def add_heading_3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_font(run, size_pt=12, bold=True, color_rgb=RGBColor(60, 60, 60))
    return p

def add_p(text, bold=False, italic=False, color_rgb=None, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_font(run, size_pt=12, bold=bold, italic=italic, color_rgb=color_rgb) # 16px = 12pt
    return p

def add_bullet(p_or_text, level=0, bold_prefix="", text=""):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        set_font(r_pre, size_pt=12, bold=True)
    if text:
        r_txt = p.add_run(text)
        set_font(r_txt, size_pt=12)
    return p

def add_image_box(img_rel_path, caption="", width_in=5.5):
    full_path = os.path.join(IMG_DIR, img_rel_path)
    if os.path.exists(full_path):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(6)
        p_img.paragraph_format.space_after = Pt(2)
        run = p_img.add_run()
        run.add_picture(full_path, width=Inches(width_in))
        if caption:
            p_cap = doc.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_cap.paragraph_format.space_after = Pt(8)
            r_cap = p_cap.add_run(f"รูปที่: {caption}")
            set_font(r_cap, size_pt=10.5, italic=True, color_rgb=RGBColor(100, 100, 100))

# ----------------- COVER PAGE -----------------
p_top = doc.add_paragraph()
p_top.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_top.paragraph_format.space_before = Pt(10)
p_top.paragraph_format.space_after = Pt(4)
r = p_top.add_run("สถาบันเทคโนโลยีพระจอมเกล้าเจ้าคุณทหารลาดกระบัง\nคณะเทคโนโลยีสารสนเทศ")
set_font(r, size_pt=16, bold=True, color_rgb=RGBColor(30, 30, 30))

add_image_box("Logo_cloudgame.png", width_in=2.0)

p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.space_before = Pt(12)
p_title.paragraph_format.space_after = Pt(6)
r_rep = p_title.add_run("รายงานโครงงานวิชา Cloud Technology (ภาคเรียนที่ 1/2569)\n")
set_font(r_rep, size_pt=15, bold=True)
r_pname = p_title.add_run("DISNEY LORCANA PLAYLAB CLOUD\n")
set_font(r_pname, size_pt=20, bold=True, color_rgb=RGBColor(202, 138, 4))
r_sub = p_title.add_run("ระบบจำลองห้องเล่นและวิเคราะห์เด็คการ์ดแบบเรียลไทม์บนคลาวด์ Serverless\n(สถานะความก้าวหน้า: เสร็จสิ้น Sprint 1–3 พร้อมส่งมอบ Stage 2)")
set_font(r_sub, size_pt=13.5, bold=False, color_rgb=RGBColor(70, 70, 70))

p_meta = doc.add_paragraph()
p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_meta.paragraph_format.space_before = Pt(20)
p_meta.paragraph_format.space_after = Pt(10)
r_m1 = p_meta.add_run("กลุ่มที่: G21\n")
set_font(r_m1, size_pt=13.5, bold=True)
r_m2 = p_meta.add_run("สาขาวิชา: เทคโนโลยีสารสนเทศ (Multimedia & Game Development / Software Engineering)\n")
set_font(r_m2, size_pt=12)
r_m3 = p_meta.add_run("สถาปัตยกรรมหลัก: 100% AWS Serverless Architecture (งบประมาณ $0.00 Free Tier)\n")
set_font(r_m3, size_pt=12, bold=True, color_rgb=RGBColor(255, 153, 0))

p_date = doc.add_paragraph()
p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_date.paragraph_format.space_before = Pt(40)
r_dt = p_date.add_run("สิงหาคม 2569")
set_font(r_dt, size_pt=12, bold=True)

doc.add_page_break()

# ----------------- ABSTRACT -----------------
add_heading_1("บทคัดย่อ (Abstract)")
add_p("โครงงาน Disney Lorcana PlayLab Cloud (กลุ่ม G21) นำเสนอการพัฒนาระบบจำลองห้องเล่นและวิเคราะห์เด็คการ์ดแบบเรียลไทม์บนสถาปัตยกรรมคลาวด์แบบไร้เซิร์ฟเวอร์ (100% AWS Serverless Architecture) สำหรับเกมการ์ดสะสม Disney Lorcana Trading Card Game (TCG) โดยมุ่งเน้นการแก้ปัญหาความหน่วงในการเชื่อมต่อ (High Latency) และภาระต้นทุนค่าใช้จ่ายเซิร์ฟเวอร์แบบเดิม (Server Overhead) ผ่านการประยุกต์ใช้บริการประมวลผลตามเหตุการณ์ (Event-Driven Computing) ด้วย AWS Lambda, ฐานข้อมูล NoSQL แบบขยายขนาดอัตโนมัติด้วย Amazon DynamoDB, และช่องทางการสื่อสารแบบสองทิศทางความเร็วสูงผ่าน AWS API Gateway WebSockets ซึ่งสามารถทำความเร็วในการซิงค์ข้อมูลสถานะกระดานระหว่างผู้เล่นได้ต่ำกว่า 100 มิลลิวินาที (Sub-100ms Latency)")

add_p("ในส่วนของส่วนติดต่อผู้ใช้ (Frontend) ระบบได้รับการพัฒนาด้วย React 19, TypeScript 5.x, Tailwind CSS v4, Framer Motion และ Zustand เพื่อมอบประสบการณ์การเล่นระดับพรีเมียม (Luxury Physical TCG Experience) ด้วยระบบฟิสิกส์ 3D การตรวจสอบการ์ดแบบสามมิติ การเปิดซองการ์ด Booster Pack เสมือนจริง และระบบจัดการเด็คการ์ดที่เชื่อมต่อกับฐานข้อมูลการ์ดอย่างเป็นทางการกว่า 408 ใบ (Set 1 The First Chapter และ Set 2 Rise of the Floodborn)")

add_p("การดำเนินงานตามระเบียบวิธีปฏิบัติงานแบบ Agile Scrum ได้ลุล่วงตามแผนงาน Sprint 1 ถึง Sprint 3 อย่างสมบูรณ์ ครอบคลุมการสร้าง Core Gameplay UI, ระบบยืนยันตัวตนและจัดการเด็ค (Authentication & Deck Microservices) และระบบห้องเล่นแบบเรียลไทม์บนคลาวด์ พร้อมทั้งผ่านการทดสอบและปรับใช้จริงบนสภาพแวดล้อม AWS Learner Lab ภายใต้นโยบายควบคุมค่าใช้จ่าย AWS Free Tier ($0.00) ได้อย่างมีประสิทธิภาพ")

# ----------------- SECTION 1 -----------------
add_heading_1("1. บทนำ (Introduction)")
add_heading_2("1.1 ความเป็นมาและความสำคัญของโครงงาน")
add_p("เกมการ์ดสะสม (Trading Card Game: TCG) เป็นหนึ่งในกิจกรรมสันทนาการและการแข่งขันระดับสากลที่มีการเติบโตอย่างต่อเนื่อง โดยเฉพาะอย่างยิ่ง Disney Lorcana TCG ที่เปิดตัวโดย Ravensburger ร่วมกับ The Walt Disney Company ซึ่งได้รับความนิยมอย่างแพร่หลายจากทั้งผู้เล่นสายการแข่งขันและนักสะสม อย่างไรก็ตาม ผู้เล่นส่วนใหญ่ยังคงเผชิญกับอุปสรรคสำคัญในการเข้าถึง ได้แก่ ต้นทุนการจัดซื้อการ์ดจริงที่มีราคาสูง การหาคู่ซ้อมที่มีระดับฝีมือใกล้เคียงกัน และข้อจำกัดด้านสถานที่ในการเล่น")

add_p("ระบบจำลองเกมการ์ดดิจิทัล (Digital TCG Simulators) ที่มีอยู่ในปัจจุบันส่วนใหญ่ประสบปัญหาสำคัญ 3 ประการ:")
add_bullet("", bold_prefix="1. ภาระต้นทุนเซิร์ฟเวอร์สูงและขยายขนาดได้ยาก: ", text="ระบบดั้งเดิมที่พึ่งพา Dedicated Server หรือ Virtual Machine (เช่น AWS EC2) ก่อให้เกิดค่าใช้จ่ายคงที่ตลอด 24 ชั่วโมงแม้ไม่มีผู้ใช้งาน และเกิดคอขวดเมื่อมีทราฟฟิกสูง")
add_bullet("", bold_prefix="2. ความหน่วงในการส่งข้อมูล (High Latency): ", text="การใช้ HTTP Polling ส่งผลให้เกิดความล่าช้าในการอัปเดตสถานะกระดานระหว่างผู้เล่น ไม่ตอบโจทย์เกมการ์ดที่ต้องการความลื่นไหลระดับมิลลิวินาที")
add_bullet("", bold_prefix="3. ประสบการณ์ผู้ใช้ที่ขาดมิติและความสมจริง: ", text="อินเทอร์เฟซส่วนใหญ่เป็นภาพกราฟิก 2 มิติที่แบนราบ ขาดมิติทางกายภาพ แอนิเมชัน และความรู้สึกของการสัมผัสการ์ดจริง")

add_p("เพื่อแก้ปัญหาเหล่านี้ โครงงาน Disney Lorcana PlayLab Cloud (กลุ่ม G21) จึงได้รับการพัฒนาขึ้นโดยใช้สถาปัตยกรรม 100% AWS Serverless Architecture ร่วมกับเทคโนโลยีเว็บสมัยใหม่ เพื่อสร้างแพลตฟอร์มจำลองห้องเล่นการ์ดและวิเคราะห์เด็คที่มีความหน่วงต่ำเป็นพิเศษ (<100ms) ปรับขนาดตามการใช้งานจริงโดยอัตโนมัติ (Elastic Auto-scaling) และควบคุมต้นทุนให้อยู่ในงบประมาณ $0.00 (AWS Free Tier) อย่างสมบูรณ์")

add_heading_2("1.2 วัตถุประสงค์ของโครงงาน")
add_bullet("", bold_prefix="1. พัฒนาระบบห้องเล่นมัลติเพลเยอร์แบบเรียลไทม์: ", text="สร้างระบบจับคู่ห้องเล่นและส่งผ่านสถานะกระดานเกมการ์ด 2 ผู้เล่นผ่าน WebSocket ที่มีความหน่วงต่ำกว่า 100ms")
add_bullet("", bold_prefix="2. สถาปัตยกรรมคลาวด์แบบ Serverless 100%: ", text="ประยุกต์ใช้ AWS Lambda, Amazon DynamoDB, AWS API Gateway และ Amazon CloudWatch รองรับการสเกลอัตโนมัติโดยไม่มีเซิร์ฟเวอร์ไอเดิล")
add_bullet("", bold_prefix="3. ยกระดับประสบการณ์ส่วนติดต่อผู้ใช้ (UX/UI): ", text="สร้างกระดานประลอง 3D Playmat, ระบบเปิดซอง Booster Gacha และระบบวิเคราะห์สมดุลเด็คการ์ด 6 ธาตุ (Ink Synergy)")
add_bullet("", bold_prefix="4. บริหารจัดการต้นทุนอย่างคุ้มค่า: ", text="ออกแบบโครงสร้างระบบให้อยู่ในกรอบ AWS Free Tier $0.00 เพื่อลดค่าใช้จ่ายในการดำเนินงานอย่างสิ้นเชิง")

# ----------------- SECTION 2 -----------------
add_heading_1("2. สถาปัตยกรรมระบบคลาวด์และเทคโนโลยี (Cloud Architecture & Tech Stack)")
add_heading_2("2.1 ผังโครงสร้างสถาปัตยกรรมระบบ (System Architecture Diagram)")
add_p("ระบบได้รับการออกแบบตามแนวทาง Event-Driven Serverless Architecture โดยแบ่งออกเป็น 3 เลเยอร์หลัก:")
add_bullet("", bold_prefix="Frontend Client: ", text="Single Page Application (SPA) พัฒนาด้วย React 19 + TypeScript + Tailwind CSS v4 โฮสต์บน Cloud Native CDN")
add_bullet("", bold_prefix="Cloud Backend Gateway: ", text="AWS API Gateway แยกออกเป็น 2 ช่องทาง ได้แก่ HTTP REST API สำหรับ CRUD เด็ค/ผู้ใช้ และ WebSocket API สำหรับ Real-Time Match Relay")
add_bullet("", bold_prefix="Serverless Compute & Database: ", text="AWS Lambda ฟังก์ชันประมวลผล (Node.js 20.x ESM) ทำงานคู่กับ Amazon DynamoDB ในการจัดเก็บ Connection ID, Match State, และ User Profiles พร้อมบันทึก Logs ไปยัง Amazon CloudWatch")

add_heading_2("2.2 เทคโนโลยีและเฟรมเวิร์กที่เลือกใช้ (Technology Stack)")
# Table for Tech Stack
tbl_tech = doc.add_table(rows=1, cols=3)
tbl_tech.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl_tech.style = 'Table Grid'
hdr_cells = tbl_tech.rows[0].cells
hdr_cells[0].text = "หมวดหมู่ (Category)"
hdr_cells[1].text = "เทคโนโลยีที่เลือกใช้ (Technology)"
hdr_cells[2].text = "บทบาทและความสำคัญ (Role & Purpose)"
for cell in hdr_cells:
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            set_font(r, size_pt=11, bold=True)

tech_data = [
    ("Frontend Core", "React 19 + TypeScript 5.x + Vite", "ส่วนติดต่อผู้ใช้แบบโมดูลาร์ ประสิทธิภาพสูง ความปลอดภัยของ Type-checking"),
    ("Styling & Anim", "Tailwind CSS v4 + Framer Motion", "การจัดเลย์เอาต์ที่ยืดหยุ่น โทนสี Luxury Dark Theme และแอนิเมชันการเล่นการ์ด"),
    ("State Management", "Zustand 5.x", "จัดการสถานะของกระดานเกม (Board State), มือการ์ด, และคลังการ์ดแบบ Reactive"),
    ("Cloud Gateway", "AWS API Gateway (REST + WebSocket)", "เราต์เตอร์รับคำขอ HTTP และจัดการการเชื่อมต่อ Full-Duplex Real-Time WebSocket"),
    ("Compute Engine", "AWS Lambda (Node.js 20.x ESM)", "ประมวลผลตรรกะเกม การตรวจสอบห้อง และการกระจายข้อความแบบ Serverless"),
    ("NoSQL Database", "Amazon DynamoDB (Pay-Per-Request)", "จัดเก็บตารางผู้ใช้, เด็คการ์ด, ห้องแข่งขัน, และ Connection Session แบบ On-demand"),
    ("Observability", "Amazon CloudWatch Logs & Metrics", "ระบบติดตามตรวจสอบประสิทธิภาพ Latency, Error Rate และการทำงานของระบบ")
]

for cat, tech, role in tech_data:
    row_cells = tbl_tech.add_row().cells
    row_cells[0].text = cat
    row_cells[1].text = tech
    row_cells[2].text = role
    for i, cell in enumerate(row_cells):
        for p in cell.paragraphs:
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.15
            for r in p.runs:
                set_font(r, size_pt=10.5, bold=(i==0))

add_p("", space_after=4)
add_heading_2("2.3 ฐานข้อมูลและการจัดเก็บข้อมูล (DynamoDB Data Models)")
add_p("โครงสร้างข้อมูลบน Amazon DynamoDB ได้รับการออกแบบให้สอดคล้องกับรูปแบบ Access Patterns ของเกมการ์ด:")
add_bullet("", bold_prefix="Lorcana_Users: ", text="Partition Key: username (String) | เก็บข้อมูลโปรไฟล์, Hash รหัสผ่าน, สถิติการชนะ/แพ้")
add_bullet("", bold_prefix="Lorcana_Decks: ", text="Partition Key: userId (String), Sort Key: deckId (String) | เก็บรายชื่อการ์ดในเด็ค (60 ใบ), ประเภท Ink, และ Metadata")
add_bullet("", bold_prefix="Lorcana_Rooms: ", text="Partition Key: roomId (String) | เก็บสถานะห้องแข่งขัน, Host ID, Guest ID, และ Turn State ปัจจุบัน")
add_bullet("", bold_prefix="Lorcana_Connections: ", text="Partition Key: connectionId (String) | เก็บ session การเชื่อมต่อ WebSocket พร้อมระบบ TTL Auto-cleanup 2 ชั่วโมง")

# ----------------- SECTION 3 -----------------
add_heading_1("3. ผลการดำเนินงานและความคืบหน้าถึง Sprint 3 (Sprint Progress & Results)")
add_heading_2("3.1 สรุปภาพรวมความก้าวหน้าราย Sprint (Sprint Breakdown)")

tbl_sprint = doc.add_table(rows=1, cols=4)
tbl_sprint.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl_sprint.style = 'Table Grid'
s_hdr = tbl_sprint.rows[0].cells
s_hdr[0].text = "Sprint"
s_hdr[1].text = "เป้าหมายหลัก (Goal)"
s_hdr[2].text = "ฟังก์ชันที่ส่งมอบ (Deliverables)"
s_hdr[3].text = "สถานะ (Status)"
for cell in s_hdr:
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            set_font(r, size_pt=11, bold=True)

sprint_data = [
    ("Sprint 1", "Core Architecture & Foundation", "ตั้งค่า React 19 Project, ออกแบบ DynamoDB Schema, พัฒนา REST API Lambda สำหรับจัดการผู้ใช้และเด็ค", "สำเร็จ 100%"),
    ("Sprint 2", "Deck Builder & Gacha Simulator", "สร้างระบบ Deck Builder 60 ใบ, Ink Synergy Analytics 6 ธาตุ, ตู้เปิดซอง Booster Pack 408 ใบ", "สำเร็จ 100%"),
    ("Sprint 3", "WebSocket Real-Time Match Board", "สร้างระบบห้องแข่งขัน 2 ผู้เล่น, กระดาน 3D Playmat, Interactive Inkwell, Action Relay Router", "สำเร็จ 100%")
]

for sp, goal, deliv, stat in sprint_data:
    r_cells = tbl_sprint.add_row().cells
    r_cells[0].text = sp
    r_cells[1].text = goal
    r_cells[2].text = deliv
    r_cells[3].text = stat
    for i, cell in enumerate(r_cells):
        for p in cell.paragraphs:
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.15
            for r in p.runs:
                set_font(r, size_pt=10.5, bold=(i==0 or i==3))

add_p("", space_after=4)
add_heading_2("3.2 รายละเอียดความสำเร็จใน Sprint 3 (Sprint 3 In-Depth Features)")
add_p("ใน Sprint ที่ 3 ทีมงานได้มุ่งเน้นการเชื่อมต่อระบบกระดานแข่งขันจริงเข้ากับ AWS WebSocket Gateway โดยมีฟังก์ชันสำคัญดังนี้:")
add_bullet("", bold_prefix="1. Real-Time 2-Player Interactive Board: ", text="กระดานประลองที่แบ่งโซนการเล่นอย่างชัดเจน (Hand, Inkwell, Play Area, Discard Pile, Lore Tracker)")
add_bullet("", bold_prefix="2. Dual-Engine Action Relay: ", text="สถาปัตยกรรมไฮบริดที่รองรับทั้งโหมด Offline Practice (Local State Engine) และโหมด Online PvP ผ่าน WebSocket Relay Router")
add_bullet("", bold_prefix="3. Inkwell Charging & Curve Control: ", text="ระบบเปลี่ยนการ์ดเป็นการ์ดหมึก (Ink) เทิร์นละ 1 ใบ พร้อมตรวจสอบเงื่อนไข Inkable/Uninkable ตามกฎมาตรฐานสากล")
add_bullet("", bold_prefix="4. Turn Transition & Lore Counting: ", text="ระบบนับแต้ม Lore สะสมสู่เป้าหมาย 20 แต้มเพื่อชัยชนะ พร้อมระบบ Ready/Set/Draw ในการเริ่มต้นเทิร์นใหม่อัตโนมัติ")

# ----------------- SECTION 4 -----------------
add_heading_1("4. ภาพประกอบหน้าจอระบบและการใช้งานจริง (Screenshots & System UI)")

add_heading_2("4.1 หน้าศูนย์รวมเกมและจับคู่ห้อง (Lobby & Match Creation)")
add_p("หน้าจอหลักสำหรับการเข้าสู่ระบบ เลือกโหมดการเล่น และสร้าง/เข้าร่วมห้องแข่งขันผ่าน Room Code:")
add_image_box("screenshot_matchlobby.png", "หน้า Match Lobby สำหรับการสร้างและเข้าร่วมห้องแข่งขันแบบเรียลไทม์", width_in=5.8)
add_image_box("screenshot_gamehub.png", "หน้า Game Hub แสดงเมนูหลักและโหมดการเล่นทั้งหมด", width_in=5.8)

add_heading_2("4.2 หน้ากระดานประลองจำลองแบบเรียลไทม์ (Interactive Playmat Battle Board)")
add_p("หน้าจอกระดานการเล่นจริงระหว่างการแข่งขัน แสดงโซนการ์ดบนมือ โซน Inkwell การ์ดในสนาม และตัวนับแต้ม Lore:")
add_image_box("screenshot_board_playing.png", "กระดานแข่งขัน 3D Playmat แสดงสถานะการเล่นและการชาร์จ Inkwell ใน Sprint 3", width_in=5.8)

add_heading_2("4.3 ระบบจัดเด็คและวิเคราะห์สมดุล (Deck Builder & Ink Synergy Analytics)")
add_p("ระบบช่วยผู้เล่นในการจัดเด็คการ์ด 60 ใบ และกราฟวิเคราะห์ค่าร่าย Ink Curve เพื่อความสมดุลสูงสุด:")
add_image_box("screenshot_deckbuilder.png", "ระบบ Deck Builder เชื่อมต่อฐานข้อมูลการ์ด 408 ใบ พร้อมตัวกรองตาม Ink Type", width_in=5.8)
add_image_box("screenshot_analytics.png", "หน้าระบบวิเคราะห์สมดุลเด็คการ์ดและสถิติ Ink Synergy", width_in=5.8)

add_heading_2("4.4 ระบบเปิดซองการ์ดเสมือนจริงและคู่มือการเล่น (Booster Gacha & Rules Guide)")
add_p("ระบบจำลองการเปิดซองการ์ดแบบ 3D และระบบคู่มือแนะนำกฎกติกาการเล่นสำหรับผู้เล่นใหม่:")
add_image_box("CardGachaDisney.png", "ระบบเปิดซอง Booster Pack เสมือนจริง พร้อมเอฟเฟกต์ความหายาก", width_in=5.8)
add_image_box("screenshot_rules.png", "หน้าคู่มือและกฎกติกาการเล่น Disney Lorcana TCG ฉบับสมบูรณ์", width_in=5.8)

# ----------------- SECTION 5 -----------------
add_heading_1("5. สรุปผลการประเมินและการดำเนินงานขั้นต่อไป (Conclusion & Next Steps)")
add_heading_2("5.1 การประเมินค่าใช้จ่ายบนคลาวด์ (Cost Assessment)")
add_p("จากการทดสอบและประมวลผลจริงบน AWS Cloud ตลอดการพัฒนา Sprint 1 ถึง Sprint 3 สรุปค่าใช้จ่ายได้ดังนี้:")
add_bullet("", bold_prefix="AWS Lambda: ", text="ใช้งานประมาณ 12,000 Invocations (ฟรี 1,000,000 Invocations/เดือน ภายใต้ Free Tier) -> $0.00")
add_bullet("", bold_prefix="Amazon DynamoDB: ", text="พื้นที่จัดเก็บ < 10 MB, Read/Write Units < 5 units (ฟรี 25 GB และ 25 WCU/RCU) -> $0.00")
add_bullet("", bold_prefix="AWS API Gateway: ", text="เชื่อมต่อ WebSocket < 5,000 นาที และคำขอ REST < 20,000 Requests -> $0.00")
add_bullet("", bold_prefix="ยอดรวมค่าใช้จ่ายทั้งหมด (Total Cost): ", text="$0.00 (สมบูรณ์ตามเป้าหมายโครงงาน)")

add_heading_2("5.2 แผนการดำเนินงานสำหรับ Sprint 4 (Stage 3 Final Sprint)")
add_bullet("", bold_prefix="1. Advanced AI Bot Opponent: ", text="พัฒนาระบบบอทปัญญาประดิษฐ์ (Decision Tree / Heuristic Engine) สำหรับโหมดฝึกซ้อมคนเดียว")
add_bullet("", bold_prefix="2. Automated Match Spectating: ", text="เพิ่มช่องทางให้ผู้ใช้งานอื่นสามารถเข้าชมการแข่งขันแบบเรียลไทม์ (Live Spectator Mode)")
add_bullet("", bold_prefix="3. Ranked Matchmaking & Leaderboards: ", text="ระบบจัดอันดับ ELO Rating และกระดานผู้นำผู้เล่นระดับสากล")
add_bullet("", bold_prefix="4. Final Production Deployment: ", text="ปรับใช้ระบบสมบูรณ์บน Custom Domain พร้อมระบบสำรองข้อมูลอัตโนมัติบน AWS S3")

doc.save(OUT_DOCX)
print(f"Successfully generated docx at {OUT_DOCX}")
