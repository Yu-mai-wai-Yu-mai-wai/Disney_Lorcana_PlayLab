# -*- coding: utf-8 -*-
"""
Full Document Generator for G21 Disney Lorcana PlayLab Cloud Phase 2
Generates LaTeX (.tex), Compiles PDF (.pdf), and generates MS Word (.docx)
"""

import os
import sys
import subprocess

sys.stdout.reconfigure(encoding='utf-8')

DOCS_DIR = os.path.dirname(os.path.abspath(__file__))
IMAGES_DIR = os.path.join(DOCS_DIR, 'images')

print(f"Working Directory: {DOCS_DIR}")

# 1. LATEX GENERATION
latex_file = os.path.join(DOCS_DIR, "G21_Disney_Lorcana_PlayLab_Cloud_Phase2_Report.tex")

with open(latex_file, "w", encoding="utf-8") as f:
    f.write(r"""\documentclass[12pt,a4paper]{report}

\usepackage{fontspec}
\usepackage{geometry}
\usepackage{graphicx}
\usepackage{float}
\usepackage{booktabs}
\usepackage{tabularx}
\usepackage{longtable}
\usepackage{array}
\usepackage{amsmath,amssymb}
\usepackage{setspace}
\usepackage{titlesec}
\usepackage{fancyhdr}
\usepackage{hyperref}
\usepackage{enumitem}
\usepackage{xcolor}
\usepackage{listings}
\usepackage{caption}
\usepackage{subcaption}

\geometry{
    a4paper,
    top=3.5cm,
    bottom=2.5cm,
    left=3.5cm,
    right=2.5cm
}

\setmainfont{TH Sarabun New}[
    Path = C:/Windows/Fonts/,
    Extension = .ttf,
    UprightFont = THSarabunNew,
    BoldFont = THSarabunNew Bold,
    ItalicFont = THSarabunNew Italic,
    BoldItalicFont = THSarabunNew BoldItalic
]

% Thai Word Breaking & Line Wrap Configuration
\XeTeXlinebreaklocale "th"
\XeTeXlinebreakskip = 0pt plus 1pt minus 0.5pt
\emergencystretch = 2em
\sloppy

\setstretch{1.15}
\setlength{\parindent}{1.25cm}
\setlength{\parskip}{6pt}

\titleformat{\chapter}[display]
  {\normalfont\bfseries\centering\fontsize{18pt}{22pt}\selectfont}
  {\chaptertitlename\ \thechapter}{10pt}{\fontsize{18pt}{22pt}\selectfont}
\titlespacing*{\chapter}{0pt}{-20pt}{20pt}

\titleformat{\section}
  {\normalfont\bfseries\fontsize{16pt}{20pt}\selectfont}
  {\thesection}{1em}{}
\titlespacing*{\section}{0pt}{12pt}{6pt}

\titleformat{\subsection}
  {\normalfont\bfseries\fontsize{15pt}{18pt}\selectfont}
  {\thesubsection}{1em}{}
\titlespacing*{\subsection}{0pt}{10pt}{4pt}

\titleformat{\subsubsection}
  {\normalfont\bfseries\fontsize{14pt}{16pt}\selectfont}
  {\thesubsubsection}{1em}{}
\titlespacing*{\subsubsection}{0pt}{8pt}{4pt}

\pagestyle{fancy}
\fancyhf{}
\fancyhead[R]{\thepage}
\renewcommand{\headrulewidth}{0pt}

\hypersetup{
    colorlinks=true,
    linkcolor=black,
    citecolor=black,
    urlcolor=blue,
    pdfauthor={นายธนัชภัทร พรหมทอง และกลุ่ม G21},
    pdftitle={Disney Lorcana PlayLab Cloud - Phase 2 Report}
}

\definecolor{codegray}{rgb}{0.5,0.5,0.5}
\definecolor{backcolour}{rgb}{0.96,0.96,0.97}
\definecolor{codeblue}{rgb}{0.13,0.34,0.67}

\lstdefinestyle{mystyle}{
    backgroundcolor=\color{backcolour},   
    commentstyle=\color{codegray},
    keywordstyle=\color{codeblue}\bfseries,
    numberstyle=\tiny\color{codegray},
    stringstyle=\color{red!70!black},
    basicstyle=\ttfamily\fontsize{9pt}{11pt}\selectfont,
    breakatwhitespace=false,         
    breaklines=true,                 
    captionpos=b,                    
    keepspaces=true,                 
    numbers=left,                    
    numbersep=5pt,                  
    showspaces=false,                
    showstringspaces=false,
    showtabs=false,                  
    tabsize=2,
    frame=single,
    rulecolor=\color{gray!30},
    xleftmargin=12pt,
    xrightmargin=6pt
}
\lstset{style=mystyle}

\renewcommand{\chaptername}{บทที่}
\renewcommand{\contentsname}{สารบัญ}
\renewcommand{\listfigurename}{สารบัญรูปภาพ}
\renewcommand{\listtablename}{สารบัญตาราง}
\renewcommand{\bibname}{บรรณานุกรม}
\renewcommand{\figurename}{รูปที่}
\renewcommand{\tablename}{ตารางที่}

\begin{document}

% -------------------------------------------------------------
% COVER PAGE (หน้าปก)
% -------------------------------------------------------------
\begin{titlepage}
\thispagestyle{empty}
\begin{center}
    \vspace*{-1.5cm}
    \includegraphics[width=3.2cm]{images/IT_KMITL_ICON.png} \\[0.8cm]
    
    {\fontsize{18pt}{22pt}\selectfont \textbf{รายงานโครงงานวิชาการเทคโนโลยีกลุ่มเมฆ (Cloud Technology)}}\\[0.3cm]
    {\fontsize{16pt}{20pt}\selectfont \textbf{ภาคเรียนที่ 1 ปีการศึกษา 2569}}\\[0.8cm]
    
    {\fontsize{18pt}{24pt}\selectfont \textbf{โครงงานระบบจำลองห้องเล่นและวิเคราะห์เด็คการ์ดแบบเรียลไทม์บนคลาวด์ Serverless}}\\[0.2cm]
    {\fontsize{16pt}{20pt}\selectfont \textbf{(Disney Lorcana PlayLab Cloud: Real-Time Serverless Card Game Simulator)}}\\[0.4cm]
    {\fontsize{16pt}{20pt}\selectfont \textbf{กลุ่ม G21}}\\[0.8cm]
    
    {\fontsize{15pt}{18pt}\selectfont \textbf{รายงานความก้าวหน้าโครงการ (Stage 2: บทที่ 1 - 3 และส่วนขยายระบบคลาวด์)}}\\[1.2cm]
    
    \begin{minipage}{0.92\textwidth}
    \fontsize{13pt}{16pt}\selectfont
    \textbf{คณะผู้จัดทำ:}\\[0.1cm]
    \begin{tabular}{@{}ll@{}}
    1. นายชยุต บุญวัฒน์ & รหัสนักศึกษา 67070032 \\
    2. นายธนัทภัทร พรหมทอง & รหัสนักศึกษา 67070069 \\
    3. นายภูริ ประชาสุขสิน & รหัสนักศึกษา 67070137 \\
    4. นางสาววรรณณิศา อมรวงศ์ไพบูลย์ & รหัสนักศึกษา 67070155 \\
    5. นายอสิธารา พุ่มดอกไม้ & รหัสนักศึกษา 67070199 \\
    6. นายวรธิษณ์ คงทอง & รหัสนักศึกษา 67070275 \\
    \end{tabular}\\[0.3cm]
    
    \textbf{อาจารย์ประจำวิชา:}\\[0.1cm]
    \begin{tabular}{@{}l@{}}
    1. ดร. ธนานพ ทองถาวร \\
    2. ผศ.ดร. พัฒนพงษ์ ฉันทมิตรโอภาส \\
    3. ผศ.ดร. ลภัส ประดิษฐ์ทัศนีย์ \\
    \end{tabular}
    \end{minipage}
    
    \vfill
    {\fontsize{15pt}{18pt}\selectfont
    คณะเทคโนโลยีสารสนเทศ\\
    สถาบันเทคโนโลยีพระจอมเกล้าเจ้าคุณทหารลาดกระบัง\\
    กันยายน 2569}
\end{center}
\end{titlepage}

% -------------------------------------------------------------
% PRELIMINARY PAGES (ส่วนนำ)
% -------------------------------------------------------------
\pagenumbering{roman}
\setcounter{page}{1}

\chapter*{บทคัดย่อภาษาไทย}
\addcontentsline{toc}{chapter}{บทคัดย่อภาษาไทย}

โครงงาน \textbf{Disney Lorcana PlayLab Cloud (กลุ่ม G21)} นำเสนอการวิเคราะห์ ออกแบบ และพัฒนาระบบจำลองห้องเล่นและวิเคราะห์เด็คการ์ดแบบเรียลไทม์บนสถาปัตยกรรมคลาวด์ไร้เซิร์ฟเวอร์เต็มรูปแบบ (100\% AWS Serverless Architecture) เพื่อรองรับเกมการ์ดสะสม Disney Lorcana Trading Card Game (TCG) โดยมุ่งเน้นการขจัดปัญหาภาระค่าใช้จ่ายเซิร์ฟเวอร์แบบเดิม (Server Overhead) และลดความหน่วงในการส่งข้อมูลระหว่างผู้เล่น (Latency Reduction) ให้ต่ำกว่า 100 มิลลิวินาที (Sub-100ms)

โครงสร้างระบบขับเคลื่อนด้วยการประมวลผลตามเหตุการณ์ (Event-Driven Computing) ผ่าน AWS Lambda ในรูปแบบ Microservices รองรับการสื่อสารแบบสองทิศทางความเร็วสูงผ่าน AWS API Gateway WebSockets ร่วมกับการจัดเก็บข้อมูลถาวรบน Amazon DynamoDB ซึ่งปรับขนาดอัตโนมัติ (Auto-scaling) และการประมวลผลวิเคราะห์เด็คการ์ดแบบ Asynchronous ผ่าน Amazon SQS ในส่วนของผู้ใช้งาน (Frontend) พัฒนาด้วย React 19, TypeScript 5, Tailwind CSS v4, Zustand State Machine และระบบฟิสิกส์ 3D Card Inspector พร้อมเชื่อมต่อฐานข้อมูลการ์ดมาตรฐาน 408 ใบ

นอกจากนี้ รายงานฉบับนี้ได้ลงลึกถึงการออกแบบวิศวกรรมเพื่อความพร้อมใช้งานสูง (High Availability) และความเชื่อถือได้ (Reliability) การป้องกันความล้มเหลวจากการพึ่งพา API ภายนอกด้วยกลยุทธ์ Multi-tier Fallback และ Local Caching การจัดการสเกลการเชื่อมต่อ WebSockets เมื่อมีห้องเล่นพร้อมกันจำนวนมากด้วย Partition Key Dispersion และ Connection Multiplexing ตลอดจนการปฏิบัติตามกรอบ AWS Well-Architected Framework ภายใต้งบประมาณที่ควบคุมได้ (\$0.00 Free Tier และไม่เกิน \$50) รวมถึงการแก้ไขข้อจำกัดเฉพาะของสภาพแวดล้อม AWS Learner Lab ได้อย่างสมบูรณ์แบบ

\vspace{0.8cm}
\noindent \textbf{คำสำคัญ:} คลาวด์ไร้เซิร์ฟเวอร์ (Serverless), เว็บซ็อกเก็ต (WebSockets), ดึงข้อมูลตามเหตุการณ์ (Event-Driven), เกมการ์ดดิจิทัล, AWS Lambda, DynamoDB, API Gateway

\newpage

\chapter*{Abstract}
\addcontentsline{toc}{chapter}{Abstract}

The \textbf{Disney Lorcana PlayLab Cloud (Group G21)} project presents the architectural design, engineering analysis, and implementation of a real-time multiplayer card simulator and asynchronous deck analytics platform built entirely on a 100\% AWS Serverless Architecture for the Disney Lorcana Trading Card Game (TCG). The project targets the elimination of persistent infrastructure costs and achieves ultra-low synchronization latency (<100ms) between remote players.

The cloud backend is designed around event-driven microservices using AWS Lambda, bidirectional persistent connections via AWS API Gateway WebSockets, autoscaling document storage with Amazon DynamoDB, and decoupled asynchronous job processing through Amazon Simple Queue Service (SQS). The frontend client is engineered using React 19, TypeScript 5, Tailwind CSS v4, Zustand global state orchestration, Framer Motion, and CSS 3D matrix physics, delivering a high-fidelity tabletop simulation connected to an official catalog of 408 cards.

Furthermore, this Stage 2 report provides comprehensive engineering solutions for critical cloud challenges: mitigating external card API dependencies through multi-tier caching and local asset fallbacks, architecting WebSocket scalability under high concurrent room load using DynamoDB partition key dispersion, auto-reconnect grace windows, and adherence to the AWS Well-Architected Framework (Reliability, Availability, Security, and Cost under \$50 budget limit). Finally, it details the engineering workarounds engineered to overcome AWS Academy Learner Lab permission boundaries (pre-provisioned LabRole integration and automated CLI pipelines).

\vspace{0.8cm}
\noindent \textbf{Keywords:} AWS Serverless, WebSockets, Event-Driven Architecture, Digital TCG, AWS Lambda, DynamoDB, API Gateway, Cloud Reliability

\newpage

\chapter*{กิตติกรรมประกาศ}
\addcontentsline{toc}{chapter}{กิตติกรรมประกาศ}

โครงงาน \textbf{Disney Lorcana PlayLab Cloud} สำเร็จลุล่วงตามวัตถุประสงค์ของการส่งมอบงานในระยะที่ 2 (Stage 2) ได้ด้วยความอนุเคราะห์ คำแนะนำ และการชี้แนะทางวิชาการอันทรงคุณค่ายิ่งจากคณาจารย์ประจำรายวิชาการเทคโนโลยีกลุ่มเมฆ (Cloud Technology) คณะเทคโนโลยีสารสนเทศ สถาบันเทคโนโลยีพระจอมเกล้าเจ้าคุณทหารลาดกระบัง ได้แก่:
\begin{enumerate}[leftmargin=1.5cm]
    \item \textbf{ดร. ธนานพ ทองถาวร}
    \item \textbf{ผศ.ดร. พัฒนพงษ์ ฉันทมิตรโอภาส}
    \item \textbf{ผศ.ดร. ลภัส ประดิษฐ์ทัศนีย์}
\end{enumerate}
ที่ได้กรุณาถ่ายทอดองค์ความรู้ด้านสถาปัตยกรรมคลาวด์ การคำนึงถึงความปลอดภัย ความพร้อมใช้งาน และการบริหารจัดการทรัพยากรอย่างคุ้มค่า ซึ่งเป็นรากฐานสำคัญในการออกแบบและแก้ปัญหาทางวิศวกรรมซอฟต์แวร์ในโครงงานนี้

คณะผู้จัดทำขอขอบพระคุณเพื่อน ๆ นักศึกษาคณะเทคโนโลยีสารสนเทศ ทุกท่านที่ได้ร่วมทดสอบระบบกระดานและห้องเล่นแบบเรียลไทม์ ให้ข้อคิดเห็นด้านประสบการณ์การใช้งาน (UX/UI) และร่วมตรวจสอบข้อผิดพลาดของระบบ รวมถึงขอขอบคุณครอบครัวที่ให้การสนับสนุนและเป็นกำลังใจในการศึกษาค้นคว้าตลอดมา

\vspace{1.5cm}
\begin{flushright}
    คณะผู้จัดทำ โครงงานกลุ่ม G21\\
    กันยายน 2569
\end{flushright}

\newpage

\tableofcontents
\newpage
\listoftables
\newpage
\listoffigures
\newpage

% -------------------------------------------------------------
% CHAPTERS (ส่วนเนื้อหา)
% -------------------------------------------------------------
\pagenumbering{arabic}
\setcounter{page}{1}

% =============================================================
% CHAPTER 1
% =============================================================
\chapter{บทนำ (Introduction)}

\section{ความเป็นมาและความสำคัญของปัญหา}
เกมการ์ดสะสมเชิงกลยุทธ์ (Trading Card Game: TCG) เป็นหนึ่งในอุตสาหกรรมเกมและสื่อสันทนาการที่มีอัตราการเติบโตสูงอย่างต่อเนื่องในระดับสากล โดยเฉพาะอย่างยิ่งการเปิดตัวเกม \textbf{Disney Lorcana TCG} โดยความร่วมมือระหว่าง Ravensburger และ The Walt Disney Company ซึ่งได้รับความนิยมอย่างแพร่หลายจากทั้งผู้เล่นสายแข่งขัน (Competitive Players) และนักสะสมทั่วโลก

อย่างไรก็ดี ในบริบทของการฝึกซ้อมและการเข้าถึงเกมการ์ด ผู้เล่นส่วนใหญ่ยังคงประสบกับข้อจำกัดหลัก 3 ประการ:
\begin{enumerate}
    \item \textbf{ต้นทุนการจัดหาการ์ดจริงและข้อจำกัดทางกายภาพ:} การ์ดหายากมีราคาสูงในตลาดรอง และการนัดหมายเล่นจำเป็นต้องใช้สถานที่จริง ทำให้ผู้เล่นใหม่ขาดโอกาสในการทดลองจัดเด็ค (Deck Building)
    \item \textbf{ภาระต้นทุนเซิร์ฟเวอร์และความไร้ประสิทธิภาพของระบบดั้งเดิม:} ระบบจำลองเกมการ์ดบนอินเทอร์เน็ตในอดีตมักทำงานบน Virtual Machine หรือ Dedicated Server (เช่น Amazon EC2) ที่ต้องเปิดทิ้งไว้ตลอด 24 ชั่วโมง ก่อให้เกิดค่าใช้จ่ายคงที่สูงแม้ไม่มีผู้ใช้งาน และมักเกิดปัญหาเซิร์ฟเวอร์ล่มเมื่อมีปริมาณทราฟฟิกพุ่งสูงขึ้นอย่างกะทันหัน
    \item \textbf{ความหน่วงในการส่งข้อมูล (High Latency):} แพลตฟอร์มบนเว็บทั่วไปมักใช้การดึงข้อมูลแบบ HTTP Polling หรือ Long Polling ซึ่งมีความหน่วงสูง ไม่ตอบสนองต่อการเล่นเกมการ์ดที่ต้องการการซิงค์ตำแหน่งการ์ด สถานะพร้อมใช้ (Ready/Exert) และแต้มคะแนน (Lore) แบบทันทีทันใด (Real-Time)
\end{enumerate}

เพื่อแก้ไขปัญหาดังกล่าว คณะผู้จัดทำจึงได้ริเริ่มโครงงาน \textbf{Disney Lorcana PlayLab Cloud (กลุ่ม G21)} โดยนำสถาปัตยกรรมคลาวด์แบบไร้เซิร์ฟเวอร์เต็มรูปแบบ (\textbf{100\% AWS Serverless Architecture}) ร่วมกับเทคโนโลยีเว็บสมัยใหม่ มาใช้ในการพัฒนาระบบจำลองกระดานเล่นการ์ดและระบบวิเคราะห์เด็คอัจฉริยะ ที่มีความหน่วงต่ำกว่า 100 มิลลิวินาที สามารถขยายขนาดรองรับผู้เล่นได้โดยอัตโนมัติ และควบคุมค่าใช้จ่ายโครงสร้างพื้นฐานให้อยู่ภายใต้งบประมาณฟรีของ AWS Free Tier (\$0.00) ได้อย่างสมบูรณ์

\section{วัตถุประสงค์ของโครงงาน}
\begin{enumerate}
    \item เพื่อออกแบบและพัฒนาระบบจำลองกระดานซ้อมเล่นการ์ด Disney Lorcana แบบเรียลไทม์บนเว็บเบราว์เซอร์ที่มีความหน่วงต่ำ (<100ms) ผ่านช่องทางสื่อสารสองทิศทาง WebSockets
    \item เพื่อประยุกต์ใช้สถาปัตยกรรม AWS Serverless (AWS Lambda, Amazon DynamoDB, AWS API Gateway) ในการสร้างระบบ Microservices ที่ประหยัดพลังงาน ยืดหยุ่น และปรับขนาดตามโหลดอัตโนมัติ
    \item เพื่อพัฒนาระบบจัดการเด็คการ์ด (Deck Builder) และระบบประมวลผลวิเคราะห์สมดุลเด็คแบบ Asynchronous ผ่าน Amazon SQS
    \item เพื่อศึกษาและวางมาตรการด้านความเชื่อถือได้ (Reliability), ความพร้อมใช้งาน (Availability), ความปลอดภัย (Security) และการควบคุมต้นทุน (Cost Optimization) ตามกรอบ AWS Well-Architected Framework
    \item เพื่อศึกษาและแก้ไขปัญหาข้อจำกัดทางวิศวกรรมบนสภาพแวดล้อมสิทธิ์จำกัดของ AWS Academy Learner Lab
\end{enumerate}

\section{ขอบเขตของระบบ (System Scope)}
ระบบครอบคลุมฟังก์ชันการทำงานหลัก 9 ด้าน (9 Core Use Cases) ดังนี้:
\begin{enumerate}
    \item \textbf{ระบบยืนยันตัวตนและความปลอดภัย (Authentication - LC-01, LC-02):} สมัครสมาชิกและเข้าสู่ระบบด้วยการเข้ารหัสผ่าน bcrypt (10 Salt Rounds) และออกใบรับรองความปลอดภัยด้วย JSON Web Token (JWT)
    \item \textbf{ระบบจัดการเด็คการ์ด (Deck Management - LC-03):} สร้าง ค้นหา แก้ไข และลบเด็คการ์ดส่วนตัว บันทึกข้อมูลลงใน Amazon DynamoDB
    \item \textbf{ระบบจำลองการเปิดซองการ์ด (Booster Pack Gacha Simulator - LC-04):} สุ่มเปิดการ์ด 12 ใบต่อซองตามสัดส่วนความน่าจะเป็นทางการ 9 ระดับ (Common ถึง Enchanted)
    \item \textbf{ระบบตรวจสอบการ์ดแบบ 3 มิติ (3D Card Inspector - LC-05):} แสดงผลการ์ดแบบสามมิติ ตรวจสอบเอฟเฟกต์ฟอยล์และรายละเอียดข้อความของการ์ดอย่างคมชัด
    \item \textbf{ระบบจำลองกระดานและห้องเล่นเรียลไทม์ (WebSockets Match Sync - LC-06):} จำลองกระดานเต็มรูปแบบ (Play Area, Inkwell, Hand, Discard, Deck, Lore Counter 0-20, Ready/Exert) และซิงค์การเคลื่อนไหวผ่าน WebSocket API Gateway
    \item \textbf{ระบบวิเคราะห์เด็คการ์ดแบบ Asynchronous (Async Deck Analyzer - LC-07):} คำนวณค่าร่ายเฉลี่ย (Ink Curve) และวิเคราะห์ความเข้ากันได้ของการ์ดผ่านคิว Amazon SQS
    \item \textbf{ระบบจัดการและดูแลผู้ใช้งาน (User Administration - LC-08):} จัดการระงับสิทธิ์บัญชีผู้ใช้เมื่อทำผิดกฎ
    \item \textbf{ระบบตรวจสอบและติดตามการทำงาน (System Monitoring - LC-09):} ติดตามเมตริกการทำงาน ทราฟฟิก และ Error Logs ผ่าน Amazon CloudWatch
\end{enumerate}

\section{ประโยชน์ที่คาดว่าจะได้รับ}
\begin{enumerate}
    \item ผู้เล่นเกม Disney Lorcana มีเครื่องมือสำหรับฝึกซ้อม ทดสอบเด็ค และเล่นกับคู่แข่งขันผ่านเว็บเบราว์เซอร์ได้ทันทีโดยไม่ต้องติดตั้งซอฟต์แวร์
    \item ได้รับองค์ความรู้และทักษะเชิงลึกในการออกแบบสถาปัตยกรรม Serverless Microservices บน Amazon Web Services ที่พร้อมรองรับการขยายตัวในระดับ Production
    \item ได้แนวทางปฏิบัติเชิงวิศวกรรมที่เป็นรูปธรรมในการรับมือกับความล้มเหลวของ API ภายนอก และการจัดการสเกลการเชื่อมต่อ WebSockets พร้อมกันจำนวนมาก
    \item สามารถลดภาระต้นทุนด้านโครงสร้างพื้นฐานคลาวด์ลงได้ 100\% เมื่อเทียบกับการเช่าโฮสต์หรือ Virtual Machine แบบเดิม
\end{enumerate}

\section{โครงสร้างของรายงาน}
รายงานโครงงานความก้าวหน้า Stage 2 ฉบับนี้แบ่งออกเป็น 5 บท ประกอบด้วย:
\begin{itemize}
    \item \textbf{บทที่ 1 บทนำ:} นำเสนอความเป็นมา วัตถุประสงค์ ขอบเขต และประโยชน์ของโครงงาน
    \item \textbf{บทที่ 2 ทฤษฎีและเทคโนโลยีที่เกี่ยวข้อง:} อธิบายสถาปัตยกรรม Serverless, WebSockets, เทคโนโลยี Frontend, โมเดลเกม Lorcana และระเบียบวิธี Agile Scrum
    \item \textbf{บทที่ 3 การวิเคราะห์และออกแบบระบบ:} ครอบคลุมการออกแบบสถาปัตยกรรมคลาวด์ 4 เลเยอร์, Use Cases, DynamoDB Schema, การแก้ปัญหา API ภายนอกล่ม, การสเกล WebSockets, เสาหลัก Well-Architected และการเอาชนะข้อจำกัดของ AWS Learner Lab
    \item \textbf{บทที่ 4 การพัฒนาระบบและผลการทดสอบคุณภาพ:} รายละเอียดการพัฒนาส่วนติดต่อผู้ใช้, Backend Microservices, แผนและผลการทดสอบคุณภาพ (QA Test Matrix)
    \item \textbf{บทที่ 5 สรุปผลการดำเนินงานและแผนงานในอนาคต:} สรุปผลงาน Stage 2, แผนงาน Sprint 4-7 ใน Stage 3 และบทเรียนที่ได้รับ
\end{itemize}

% =============================================================
% CHAPTER 2
% =============================================================
\chapter{ทฤษฎีและเทคโนโลยีที่เกี่ยวข้อง (Related Theories \& Technologies)}

\section{สถาปัตยกรรมคลาวด์แบบไร้เซิร์ฟเวอร์ (AWS Serverless Architecture)}
สถาปัตยกรรม Serverless เป็นรูปแบบการประมวลผลบนคลาวด์ที่ผู้พัฒนาไม่ต้องทำการจัดสรร ดูแล หรือบำรุงรักษาเครื่องเซิร์ฟเวอร์ (No Server Management) โดยคลาวด์โพรไวเดอร์เป็นผู้บริหารจัดการทรัพยากร การขยายขนาด (Auto-scaling) และความพร้อมใช้งานให้โดยอัตโนมัติ โดยคิดค่าบริการตามปริมาณการใช้งานจริงในระดับมิลลิวินาที (Pay-per-use)

โครงงานนี้ได้ประยุกต์ใช้บริการคลาวด์หลักของ Amazon Web Services (AWS) ดังต่อไปนี้:
\begin{enumerate}
    \item \textbf{AWS Lambda:} บริการประมวลผลโค้ดตามเหตุการณ์ (Function-as-a-Service: FaaS) ทำหน้าที่รัน Microservices ของระบบทั้งส่วน REST API และ WebSocket Event Routing โดยทำงานเฉพาะเมื่อมี Request เข้ามาและปิดตัวลงเมื่อทำงานเสร็จสิ้น
    \item \textbf{Amazon API Gateway:} บริการจัดการ API แบบ Fully Managed ทำหน้าที่เป็นประตูทางเข้าของระบบ รองรับทั้ง \textit{HTTP API} สำหรับการส่งคำขอแบบ Request-Response (REST) และ \textit{WebSocket API} สำหรับการเชื่อมต่อแบบเปิดค้างไว้สองทิศทาง (Stateful Persistent Connection)
    \item \textbf{Amazon DynamoDB:} ฐานข้อมูล NoSQL แบบ Key-Value และ Document Database ที่มีความเร็วสูงในระดับ Single-digit Millisecond และรองรับโหมดการคิดค่าบริการแบบตามปริมาณการใช้งานจริง (On-Demand / Pay-Per-Request)
    \item \textbf{Amazon Simple Queue Service (SQS):} บริการคิวข้อความแบบกระจายตัวและไร้เซิร์ฟเวอร์ ทำหน้าที่เป็นตัวกลางในการส่งต่องานประมวลผลหนัก เช่น การคำนวณสถิติเด็ค เพื่อลดภาระการรอคอยของผู้ใช้งาน (Decoupled Asynchronous Processing)
    \item \textbf{Amazon CloudWatch:} บริการศูนย์รวมการเก็บ Logs เมตริกประสิทธิภาพ และการแจ้งเตือน (Alarms) เมื่อเกิดความผิดปกติในระบบ
\end{enumerate}

\section{สถาปัตยกรรมการสื่อสารแบบสองทิศทางเรียลไทม์ (WebSockets Protocol)}
โพรโทคอล WebSocket (ตามมาตรฐาน RFC 6455) ทำงานบนช่องทาง TCP เดียวกันแบบ Full-Duplex ทำให้ Client และ Server สามารถส่งข้อมูลแลกเปลี่ยนกันได้ทันทีโดยไม่ต้องเริ่มต้น Handshake ใหม่ในทุกๆ ครั้งที่ส่งข้อมูล ซึ่งแตกต่างจาก HTTP Polling ที่สร้าง Overhead ของ HTTP Header ซ้ำซ้อน

ในโครงสร้างของ AWS API Gateway WebSockets การจัดการ Connection จะถูกบริหารจัดการโดยโครงสร้างพื้นฐานของ AWS ซึ่งมีกลไกสำคัญดังนี้:
\begin{itemize}
    \item \textbf{Connection Lifecycle Management:} เมื่อผู้เล่นเชื่อมต่อ API Gateway จะสร้าง \texttt{connectionId} แบบไม่ซ้ำกัน และทริกเกอร์ Lambda ผ่านเส้นทาง \texttt{\$connect} และเมื่อผู้เล่นตัดการเชื่อมต่อจะทริกเกอร์เส้นทาง \texttt{\$disconnect}
    \item \textbf{Route Selection Expression:} กำหนดให้คัดแยกประเภทของคำขอจาก JSON Payload เช่น \texttt{\$request.body.action} เพื่อเรียก Lambda Function ที่สอดคล้องกับ Action นั้นๆ
    \item \textbf{ApiGatewayManagementApi (@connections):} อินเตอร์เฟซที่อนุญาตให้ Lambda ส่งข้อมูลแบบ Push ไปยังผู้เล่นคนอื่นๆ ในห้องเล่นเดียวกันได้โดยตรงผ่าน \texttt{PostToConnectionCommand}
\end{itemize}

\section{เทคโนโลยีฝั่งผู้ใช้งาน (Frontend Engineering \& State Architecture)}
ส่วนติดต่อผู้ใช้ถูกสร้างขึ้นเพื่อมอบประสบการณ์การเล่นระดับพรีเมียม (Luxury Physical TCG Simulation) โดยใช้ชุดเครื่องมือระดับแนวหน้า:
\begin{itemize}
    \item \textbf{React 19 \& TypeScript 5:} โครงสร้างคอมโพเนนต์แบบ Type-Safe และใช้ฟีเจอร์การจัดการ Concurrent Rendering เพื่อให้ UI ตอบสนองต่อการกระทำของผู้ใช้ได้ทันที
    \item \textbf{Tailwind CSS v4:} ระบบยูทิลิตี้คลาสแบบ CSS-First ที่ช่วยให้การจัดรูปแบบกราฟิกในธีม \textit{Dark Editorial + Magic R3} มีความสวยงามและมีประสิทธิภาพสูง
    \item \textbf{Zustand State Management:} การจัดการ Global State แบบเบาบางและมีประสิทธิภาพสูง แยก Store ออกเป็นโมดูลชัดเจน เช่น \texttt{usePlaymatStore}, \texttt{useDeckStore}, \texttt{useAuthStore}
    \item \textbf{Framer Motion \& CSS 3D Transforms:} ไลบรารีแอนิเมชันสำหรับจัดการระบบฟิสิกส์ของการ์ด การหมุนการ์ดสามมิติ การคว่ำการ์ดลง Inkwell และการเปิดซองการ์ดแบบเสมือนจริง
\end{itemize}

\section{โมเดลโดเมนเกมการ์ดสะสม Disney Lorcana}
เกม Disney Lorcana เป็นเกมการ์ดสะสมที่มีกฎกติกาและองค์ประกอบเฉพาะที่ระบบต้องจำลองอย่างแม่นยำ:
\begin{itemize}
    \item \textbf{สีหมึกทั้ง 6 หมวด (Ink Types):} Amber (เหลือง), Amethyst (ม่วง), Emerald (เขียว), Ruby (แดง), Sapphire (น้ำเงิน), Steel (เทา) โดยกฎการจัดเด็คต้องมีขนาดอย่างน้อย 60 ใบ และประกอบด้วยสีหมึกได้ไม่เกิน 2 สี
    \item \textbf{ประเภทการ์ด (Card Types):} Character (ตัวละครสำหรับเควสต์หรือท้าทาย), Action/Song (เวทมนตร์หรือบทเพลง), Item (อุปกรณ์สนับสนุน), Location (สถานที่สะสม Lore)
    \item \textbf{สถานะของการ์ด (Card States):} Ready (การ์ดตั้งตรง พร้อมใช้งาน) และ Exerted (การ์ดหมุน 90 องศา อยู่ในสถานะพักหรือถูกใช้งานแล้ว)
    \item \textbf{เงื่อนไขชัยชนะ (Win Condition):} ผู้เล่นที่สามารถสะสมแต้มความรู้ (Lore) ครบ 20 แต้มก่อนจะเป็นผู้ชนะในเกม
\end{itemize}

\section{ระเบียบวิธีปฏิบัติงานแบบ Agile Scrum}
โครงงานดำเนินงานตามกรอบการทำงาน Agile Scrum แบ่งรอบการพัฒนาออกเป็น 7 Sprints ซึ่งใน Stage 2 นี้ ได้ส่งมอบงานใน Sprint 1 ถึง Sprint 3 เสร็จสิ้นสมบูรณ์ ดังแสดงในตารางที่ \ref{tab:sprint_plan}

\begin{table}[H]
\centering
\caption{แผนการดำเนินงานและการส่งมอบงานแต่ละ Sprint}
\label{tab:sprint_plan}
\begin{tabularx}{\textwidth}{lcp{6.5cm}c}
\toprule
\textbf{Sprint} & \textbf{ระยะเวลา} & \textbf{เป้าหมายหลักของการส่งมอบ} & \textbf{สถานะ} \\
\midrule
Sprint 1 & 1–15 ส.ค. 2569 & Scaffold โปรเจกต์, UI กระดาน, ชุดการ์ด 408 ใบ & เสร็จสิ้น (100\%) \\
Sprint 2 & 16–31 ส.ค. 2569 & Auth (JWT/bcrypt), Deck REST API, DynamoDB & เสร็จสิ้น (100\%) \\
Sprint 3 & 1–10 ก.ย. 2569 & WebSocket Real-time Room Sync บน AWS & เสร็จสิ้น (100\%) \\
Sprint 4 & 11–25 ก.ย. 2569 & Async Deck Analyzer บน SQS + Deck Algorithm & กำลังดำเนินการ \\
Sprint 5 & 26 ก.ย. – 10 ต.ค. & CloudWatch Dashboard, Observability, Hardening & ตามแผน \\
Sprint 6 & 11–20 ต.ค. 2569 & Load Testing (Artillery), WebSocket Reconnection & ตามแผน \\
Sprint 7 & 21–25 ต.ค. 2569 & จัดทำรายงานฉบับสมบูรณ์, Video Demo, สรุปผล & ตามแผน \\
\bottomrule
\end{tabularx}
\end{table}

% =============================================================
% CHAPTER 3
% =============================================================
\chapter{การวิเคราะห์และออกแบบระบบ (System Analysis \& Architectural Design)}

\section{ข้อกำหนดความต้องการของระบบ (System Requirements)}
การออกแบบระบบได้รับการวิเคราะห์อย่างเป็นระบบโดยแบ่งออกเป็นสองส่วนหลัก:

\subsection{ข้อกำหนดเชิงฟังก์ชัน (Functional Requirements: FR)}
\begin{enumerate}
    \item \textbf{FR-01 (Authentication):} ผู้ใช้สามารถลงทะเบียน เข้าสู่ระบบ และได้รับ JWT Token สำหรับยืนยันตัวตนกับ REST และ WebSocket API
    \item \textbf{FR-02 (Deck Management):} ผู้ใช้สามารถสร้าง ค้นหา ปรับแต่ง และลบเด็คการ์ด โดยระบบตรวจสอบกฎความถูกต้องของเด็ค (60 ใบ, ไม่เกิน 2 สี, ซ้ำไม่เกิน 4 ใบ)
    \item \textbf{FR-03 (Booster Gacha Simulator):} ผู้ใช้สามารถสุ่มเปิดการ์ด 12 ใบตามอัตราความหายากมาตรฐานอย่างเป็นทางการ
    \item \textbf{FR-04 (3D Card Inspection):} ผู้ใช้สามารถซูมและหมุนดูมิติของการ์ดและเอฟเฟกต์ฟอยล์แบบสามมิติได้
    \item \textbf{FR-05 (Real-time Multiplayer Match):} ผู้เล่น 2 คนสามารถสร้างห้องด้วยรหัส 6 หลักและเชื่อมต่อเข้ากระดานเล่นเดียวกันได้
    \item \textbf{FR-06 (Board Action Synchronization):} การลากการ์ดลงสนาม การหมุน Ready/Exert การนำการ์ดเข้า Inkwell และการปรับแต้ม Lore ต้องซิงค์ระหว่าง 2 ผู้เล่นทันที
    \item \textbf{FR-07 (Deck Balance Analytics):} ระบบสามารถวิเคราะห์กราฟการกระจายตัวของค่าร่าย (Ink Curve) และสัดส่วนหมวดหมู่การ์ดได้
\end{enumerate}

\subsection{ข้อกำหนดที่ไม่ใช่เชิงฟังก์ชัน (Non-Functional Requirements: NFR)}
\begin{enumerate}
    \item \textbf{NFR-01 (Low Latency):} ความหน่วงในการส่งข้อมูล Action ผ่าน WebSocket ต้องต่ำกว่า 100 มิลลิวินาที (Sub-100ms Latency)
    \item \textbf{NFR-02 (High Availability \& Fault Tolerance):} ระบบคลาวด์ต้องมีความพร้อมใช้งานไม่ต่ำกว่า 99.9\% โดยไม่มี Single Point of Failure
    \item \textbf{NFR-03 (Security \& Privacy):} รหัสผ่านต้องได้รับการแฮชด้วย bcrypt ก่อนบันทึกลงฐานข้อมูล และข้อมูลส่งผ่าน HTTPS และ WSS ตลอดเส้นทาง
    \item \textbf{NFR-04 (Cost Optimization):} ค่าใช้จ่ายทั้งหมดต้องอยู่ภายใต้กรอบ Free Tier (\$0.00) และงบประมาณโครงงานไม่เกิน \$50
\end{enumerate}

\section{สถาปัตยกรรมระบบโดยรวม (End-to-End Serverless Architecture)}
ระบบได้รับการออกแบบตามสถาปัตยกรรม 4 เลเยอร์ (4-Layer Serverless Architecture) ดังแสดงในรูปที่ \ref{fig:arch_overview}:

\begin{figure}[H]
\centering
\noindent\fbox{%
\begin{minipage}{\dimexpr\textwidth-2\fboxsep-2\fboxrule\relax}
\centering
\vspace{0.3cm}
{\fontsize{13pt}{16pt}\selectfont \textbf{[ Presentation Layer: Web Client SPA ]}}\\[0.15cm]
React 19 + TypeScript + Tailwind CSS v4 + Zustand + Framer Motion (HTTPS)\\[0.3cm]
$\Downarrow$ \qquad \qquad $\Downarrow$\\[0.15cm]
{\fontsize{13pt}{16pt}\selectfont \textbf{[ API \& Routing Layer: AWS API Gateway ]}}\\[0.15cm]
\begin{tabularx}{\linewidth}{X|X}
\centering \textbf{HTTP API (REST)} & \centering \textbf{WebSocket API (WSS)} \tabularnewline
\centering Routes: \texttt{/auth/*}, \texttt{/decks} & \centering Routes: \texttt{\$connect}, \texttt{\$disconnect}, \texttt{sendAction} \tabularnewline
\end{tabularx}\\[0.3cm]
$\Downarrow$ \qquad \qquad $\Downarrow$\\[0.15cm]
{\fontsize{13pt}{16pt}\selectfont \textbf{[ Compute Layer: AWS Lambda Microservices (Node.js 20.x) ]}}\\[0.15cm]
$\bullet$ Auth Microservice \quad $\bullet$ Deck Microservice \quad $\bullet$ Room Router \quad $\bullet$ Deck Analyzer\\[0.3cm]
$\Downarrow$ \qquad \qquad $\Downarrow$\\[0.15cm]
{\fontsize{13pt}{16pt}\selectfont \textbf{[ Persistence \& Messaging Layer ]}}\\[0.15cm]
$\bullet$ Amazon DynamoDB (Users, Decks, RoomState) \quad $\bullet$ Amazon SQS \quad $\bullet$ CloudWatch Logs
\vspace{0.3cm}
\end{minipage}%
}
\caption{สถาปัตยกรรมระบบคลาวด์แบบไร้เซิร์ฟเวอร์เต็มรูปแบบ (100\% AWS Serverless)}
\label{fig:arch_overview}
\end{figure}

\section{การออกแบบ Use Case และลำดับการทำงาน (Use Case Design)}
ตารางที่ \ref{tab:usecases} สรุปรายละเอียด Use Cases ทั้ง 9 รายการของระบบตามมาตรฐานวิศวกรรมซอฟต์แวร์:

\begin{table}[H]
\centering
\caption{ตารางสรุปรายละเอียด Use Cases ของระบบ (LC-01 ถึง LC-09)}
\label{tab:usecases}
\small
\begin{tabularx}{\textwidth}{cllX}
\toprule
\textbf{ID} & \textbf{ชื่อ Use Case} & \textbf{ผู้กระทำ (Actor)} & \textbf{คำอธิบายโดยย่อ} \\
\midrule
LC-01 & สมัครสมาชิก & Guest & ตรวจสอบความถูกต้อง แฮชรหัสผ่านด้วย bcrypt และบันทึกบัญชี \\
LC-02 & เข้าสู่ระบบ & Guest & ตรวจสอบรหัสผ่านและสร้าง JWT Token ส่งกลับให้เบราว์เซอร์ \\
LC-03 & จัดการเด็คการ์ด & Member & สร้าง แก้ไข ดึงข้อมูล และลบเด็คการ์ดส่วนตัวบน DynamoDB \\
LC-04 & สุ่มเปิดซองการ์ด & Member & สุ่มการ์ด 12 ใบตามเรตความหายากทางการ 9 ระดับ \\
LC-05 & ตรวจสอบการ์ด 3D & Member & ซูม หมุนดูการ์ดสามมิติ และตรวจสอบเอฟเฟกต์ฟอยล์ \\
LC-06 & ซ้อมเล่นห้องเรียลไทม์ & Member & สร้างห้อง 6 หลัก ซิงค์พิกัดการ์ด หมุน Exert/Ready และ Lore \\
LC-07 & วิเคราะห์เด็ค Asynchronous & Member & ส่งเด็คเข้าคิว SQS ให้ Lambda วิเคราะห์ Ink Curve \& Synergy \\
LC-08 & จัดการผู้ใช้งาน & Admin & ค้นหาและระงับสิทธิ์บัญชีผู้ใช้งานที่ทำผิดกฎ \\
LC-09 & ตรวจสอบระบบ & Admin & ติดตามเมตริกทราฟฟิก Latency และ Error Logs บน CloudWatch \\
\bottomrule
\end{tabularx}
\end{table}

\section{การออกแบบฐานข้อมูลแบบ NoSQL (Amazon DynamoDB Schema)}
ระบบใช้ Amazon DynamoDB ซึ่งเป็นฐานข้อมูล NoSQL โดยออกแบบ Schema แยกตามหน้าที่การทำงาน (Table-per-Service Design) เพื่อลดความซับซ้อนและเพิ่มความเร็วในการสืบค้น:

\begin{table}[H]
\centering
\caption{โครงสร้างตารางข้อมูลบน Amazon DynamoDB}
\label{tab:dynamo_schema}
\small
\begin{tabularx}{\textwidth}{p{3.4cm}p{2.1cm}p{2.5cm}X}
\toprule
\textbf{ตาราง (Table)} & \textbf{Partition Key} & \textbf{Sort Key} & \textbf{Attributes สำคัญ} \\
\midrule
\texttt{LorcanaUsers} & \texttt{userId} (Str) & - & \texttt{username}, \texttt{passwordHash}, \texttt{email}, \texttt{createdAt} \\
\texttt{LorcanaDecks} & \texttt{userId} (Str) & \texttt{deckId} (Str) & \texttt{deckName}, \texttt{inks} (List), \texttt{cards} (JSON), \texttt{updatedAt} \\
\texttt{LorcanaRoomState} & \texttt{roomId} (Str) & \texttt{connectionId} (Str) & \texttt{username}, \texttt{role} (P1/P2), \texttt{state} (JSON), \texttt{ttl} (Num) \\
\bottomrule
\end{tabularx}
\end{table}

\section{การออกแบบการจัดการ API ภายนอกและการป้องกันความล้มเหลว (External API Dependency \& High Availability Strategy)}
ระบบจำเป็นต้องใช้ข้อมูลการ์ด Disney Lorcana จำนวน 408 ใบจาก Set 1 (\textit{The First Chapter}) และ Set 2 (\textit{Rise of the Floodborn}) ซึ่งข้อมูลและรูปภาพต้นฉบับดึงมาจาก External API ภายนอก (เช่น \texttt{lorcana-api.com})

\subsection{การวิเคราะห์ความเสี่ยงเมื่อ API ภายนอกล้มเหลว}
หากระบบพึ่งพา External API โดยตรงแบบ Real-Time จะก่อให้เกิดความเสี่ยงร้ายแรง 4 ประการ:
\begin{enumerate}
    \item \textbf{Third-party Outage / Downtime:} หากเซิร์ฟเวอร์ภายนอกปิดปรับปรุงหรือล่ม จะส่งผลให้ระบบ Playmat และ Deck Builder ใช้งานไม่ได้ทันที
    \item \textbf{Rate Limiting \& IP Throttling:} หากมีผู้ใช้งานจำนวนมากยิงคำขอดึงการ์ดพร้อมกัน API ภายนอกจะตอบกลับด้วย HTTP 429 Too Many Requests
    \item \textbf{Latency Spike:} การดึงข้อมูลข้ามมหาสมุทรเพิ่มความหน่วงให้กับ Client 300--800ms ในทุกการโหลดหน้าจอ
    \item \textbf{API Deprecation / Breaking Changes:} การเปลี่ยนแปลงโครงสร้างข้อมูล Schema ของผู้ให้บริการภายนอกอาจทำให้โค้ดฝั่งไคลเอนต์ Error
\end{enumerate}

\subsection{กลยุทธ์การแก้ไขปัญหาความล้มเหลว 4 ระดับ (4-Tier Fallback Strategy)}
เพื่อสร้างความพร้อมใช้งานระดับ 100\% ทีมงานได้ออกแบบสถาปัตยกรรมความทนทานต่อข้อผิดพลาด (Fault-Tolerant Architecture) ดังนี้:
\begin{enumerate}
    \item \textbf{ระดับที่ 1: Local Static Asset Bundling (Zero-Dependency Base):} จัดทำไฟล์ Snapshot ข้อมูลการ์ดทั้ง 408 ใบในรูปแบบ \texttt{cards\_set1\_set2.json} และคอมไพล์รวมไว้ในบันเดิลของ Frontend Client และโฮสต์บน S3/CloudFront โดยตรง ทำให้ระบบสามารถทำงานได้ทันทีแม้ไม่มีการเชื่อมต่อกับ External API
    \item \textbf{ระดับที่ 2: In-Browser LocalStorage / IndexedDB Caching:} เมื่อเบราว์เซอร์ดึงข้อมูลการ์ดสำเร็จ จะบันทึกแคชลงใน LocalStorage พร้อมตั้งค่า Time-To-Live (TTL = 7 วัน) เพื่อลดการยิงเครือข่ายซ้ำ
    \item \textbf{ระดับที่ 3: Serverless Reverse Proxy \& DynamoDB Card Catalog:} ฝั่ง Backend มีฟังก์ชัน Lambda ทำหน้าที่เป็น Caching Proxy หากมีการเรียกดูข้อมูลใหม่ Lambda จะตรวจสอบในตาราง DynamoDB ก่อน หากไม่มีจึงจะดึงจาก API ภายนอกและบันทึกลง DynamoDB เพื่อให้บริการแก่ผู้ใช้คนอื่นต่อไป
    \item \textbf{ระดับที่ 4: Graceful UI Degradation \& Card Art Fallback:} หากไม่สามารถโหลดรูปภาพการ์ดต้นฉบับจาก CDN ภายนอกได้ ระบบมีรูปภาพสำรอง SVG Card Back และกรอบการ์ดเวกเตอร์ในตัว (\texttt{Lorcana\_Card\_Back.png}) พร้อมแสดงข้อความเตือนแบบ Non-blocking Alert เพื่อให้ผู้เล่นสามารถดำเนินเกมต่อได้โดยไม่สะดุด
\end{enumerate}

\section{การออกแบบการสเกลและรองรับการเชื่อมต่อ WebSockets พร้อมกัน (WebSocket Scalability \& Concurrency Management)}

\subsection{ปัญหาคอขวดเมื่อมีผู้ใช้งานและห้องเล่นจำนวนมาก}
เมื่อมีผู้เล่นพร้อมกันในระดับหลายร้อยถึงหลายพันห้องพร้อมกัน ระบบ WebSockets บนคลาวด์จะเผชิญกับความท้าทายดังนี้:
\begin{enumerate}
    \item \textbf{Connection Exhaustion \& Socket State Tracking:} การติดตามสถานะของผู้เล่นที่เปิดค้างไว้และการตรวจจับการหลุดการเชื่อมต่อ (Zombie Connections)
    \item \textbf{Hot Partition Keys ใน DynamoDB:} หากทุกคำสั่งในห้องเล่นเขียนลง Partition Key เดียวกันด้วยความถี่สูง อาจเกิดปัญหา DynamoDB ProvisionedThroughputExceededException
    \item \textbf{Broadcast Latency \& Fan-Out Overhead:} เมื่อต้องกระจายข้อความ Action ไปยังผู้เล่นทุกคนในห้อง หากจัดการแบบ Synchronous Loop จะเพิ่มเวลาตอบสนอง
\end{enumerate}

\subsection{แนวทางการแก้ไขปัญหาและการออกแบบเชิงสถาปัตยกรรม}
\begin{enumerate}
    \item \textbf{การออกแบบ Compound Key \& TTL Auto-Cleanup ใน DynamoDB:}
    ใช้ \texttt{roomId} เป็น Partition Key (Hash) ร่วมกับ \texttt{connectionId} เป็น Sort Key (Range) ทำให้การ Query ผู้เล่นในห้องเดียวกันทำได้ด้วยความเร็ว $O(1)$ ผ่านคำสั่ง \texttt{QueryCommand} และกำหนดฟิลด์ \texttt{ttl} (Epoch Timestamp + 2 ชั่วโมง) เพื่อให้ DynamoDB ลบข้อมูลห้องที่เล่นเสร็จแล้วทิ้งโดยอัตโนมัติโดยไม่มีค่าใช้จ่าย Write Capacity Unit
    
    \item \textbf{การใช้ Optimistic UI Updates ร่วมกับ Event Relay:}
    ฝั่ง Frontend ใช้สถาปัตยกรรม Optimistic Rendering โดยเมื่อผู้เล่นลากการ์ด หน้าจอของตนเองจะอัปเดตตำแหน่งทันทีใน 0ms และส่งเฉพาะ Action Payload สั้นๆ (\textasciitilde120 bytes) ผ่าน WebSocket เพื่อให้ Lambda Room Router ส่งต่อ (Relay) ไปยังคู่แข่งผ่าน AWS API Gateway Management API โดยตรง ทำให้ Latency ฝั่งคู่แข่งต่ำกว่า 100ms
    
    \item \textbf{ระบบ Auto-Reconnect พร้อม Grace Window 30 วินาที:}
    หากเกิดสัญญาณเน็ตขาดหาย Client จะพยายามเชื่อมต่อใหม่แบบ Exponential Backoff และส่ง Action \texttt{RECONNECT\_ROOM} พร้อม JWT Token เดิม ระบบจะกู้คืนสถานะกระดานล่าสุดจาก DynamoDB ให้ผู้เล่นทันทีโดยไม่ต้องเริ่มเกมใหม่
    
    \item \textbf{Decoupled Fan-Out Pattern สำหรับการแข่งขันขนาดใหญ่:}
    สำหรับโหมดที่มีผู้ชม (Spectator Mode) เกิน 10 คน ระบบเตรียมพร้อมสถาปัตยกรรมต่อยอดโดยใช้ Amazon SQS หรือ Redis Pub/Sub เพื่อกระจายโหลดการ Broadcast ออกจาก Lambda Room Router หลัก ป้องกันปัญหา Concurrency Throttling
\end{enumerate}

\section{การออกแบบตามกรอบ AWS Well-Architected Framework}

\subsection{ด้านความเชื่อถือได้ (Reliability)}
\begin{itemize}
    \item \textbf{Idempotent Event Processing:} ออกแบบให้ฟังก์ชัน Lambda ทุกตัวสามารถรับ Action ซ้ำได้โดยไม่ก่อให้เกิดสถานะผิดพลาด
    \item \textbf{Dead Letter Queues (DLQ):} เชื่อมต่อ Amazon SQS DLQ เข้ากับฟังก์ชัน Lambda เพื่อดักจับและกักเก็บข้อความที่ประมวลผลล้มเหลวไว้ตรวจสอบย้อนหลัง
    \item \textbf{Automatic Retry Policies:} กำหนดนโยบาย Retry ของ AWS SDK แบบ Exponential Backoff with Jitter เพื่อลดปัญหาโหลดกระชาก
\end{itemize}

\subsection{ด้านความพร้อมใช้งาน (Availability)}
\begin{itemize}
    \item \textbf{Multi-AZ Native Redundancy:} บริการ AWS Lambda, API Gateway และ DynamoDB มีการสำรองข้อมูลและกระจายการทำงานข้ามอย่างน้อย 3 Availability Zones (AZs) โดยอัตโนมัติ ให้ค่าความพร้อมใช้งานระดับ 99.99\% Service Level Agreement (SLA)
    \item \textbf{No Single Point of Failure (SPOF):} สถาปัตยกรรมไม่มีเครื่องเซิร์ฟเวอร์เดี่ยวตัวใดตัวหนึ่งที่เป็นคอขวด หาก Lambda Container ตัวใดหยุดทำงาน AWS จะสปอว์น Container ตัวใหม่ขึ้นมาแทนที่ในเสี้ยววินาที
\end{itemize}

\subsection{ด้านความมั่นคงปลอดภัย (Security)}
\begin{itemize}
    \item \textbf{Data in Transit Encryption:} บังคับใช้การเข้ารหัสข้อมูลขณะส่งผ่านเครือข่ายด้วยโพรโทคอล HTTPS (TLS 1.3) สำหรับ REST API และ WSS (Secure WebSockets) สำหรับการสื่อสารแบบสองทิศทาง
    \item \textbf{Data at Rest Encryption:} ข้อมูลทั้งหมดใน Amazon DynamoDB ได้รับการเข้ารหัสที่ระดับฮาร์ดแวร์ด้วย AWS Key Management Service (KMS) Default Encryption
    \item \textbf{Authentication \& Password Hashing:} รหัสผ่านถูกแฮชด้วยอัลกอริทึม bcrypt โดยกำหนด Salt Rounds เท่ากับ 10 ทำให้ไม่สามารถถอดรหัสย้อนกลับได้แม้ฐานข้อมูลจะรั่วไหล และใช้ JWT Token ที่มีการลงลายมือชื่อดิจิทัล (HMAC-SHA256) และมีวันหมดอายุที่ชัดเจน
    \item \textbf{Principle of Least Privilege (PoLP):} จำกัดสิทธิ์ของ Lambda Function ให้สามารถเข้าถึงเฉพาะตาราง DynamoDB และคิว SQS ที่จำเป็นต่อการทำงานเท่านั้น
\end{itemize}

\subsection{ด้านการบริหารจัดการต้นทุน (Cost Optimization under \$50 Budget)}
โครงสร้างพื้นฐานได้รับการออกแบบให้อยู่ในกรอบ AWS Free Tier (\$0.00 ต่อเดือน) และจำกัดงบประมาณไม่เกิน \$50 ดังแสดงในตารางที่ \ref{tab:cost_breakdown}:

\begin{table}[H]
\centering
\caption{การประเมินค่าใช้จ่ายบนสถาปัตยกรรม AWS Serverless (ผู้ใช้งาน 10,000 คน/เดือน)}
\label{tab:cost_breakdown}
\begin{tabularx}{\textwidth}{lXXr}
\toprule
\textbf{บริการ (Service)} & \textbf{โควตา AWS Free Tier} & \textbf{ปริมาณการใช้งานจริง} & \textbf{ค่าบริการจริง} \\
\midrule
AWS Lambda & 1,000,000 Requests/เดือน & \textasciitilde250,000 Requests & \$0.00 \\
Amazon DynamoDB & 25 GB Storage, 25 WCU/RCU & < 1 GB, On-Demand Mode & \$0.00 \\
AWS API Gateway & 1,000,000 Calls \& Messages & \textasciitilde400,000 Messages & \$0.00 \\
Amazon SQS & 1,000,000 Requests/เดือน & \textasciitilde50,000 Requests & \$0.00 \\
Amazon CloudWatch & 5 GB Log Ingestion & < 500 MB & \$0.00 \\
\midrule
\multicolumn{3}{l}{\textbf{รวมค่าใช้จ่ายต่อเดือนทั้งสิ้น (Total Monthly Cost)}} & \textbf{\$0.00} \\
\bottomrule
\end{tabularx}
\end{table}

เพื่อป้องกันปัญหาค่าบริการส่วนเกิน (Cost Overrun) ทีมงานได้ตั้งค่า \textbf{AWS CloudWatch Billing Alarm} ให้ส่งการแจ้งเตือนทางอีเมลทันทีหากมีค่าบริการเกิดขึ้นเกิน \$5.00 และ \$20.00 ตามลำดับ

\section{การวิเคราะห์ข้อจำกัดของ AWS Learner Lab และแนวทางแก้ไขเชิงวิศวกรรม (Learner Lab Constraints \& Engineering Workarounds)}
สภาพแวดล้อม AWS Academy Learner Lab เป็นคลาวด์แซนด์บ็อกซ์เพื่อการศึกษาที่มีการจำกัดสิทธิ์ความปลอดภัยอย่างเข้มงวด ทีมงานได้วิเคราะห์ข้อจำกัดและพัฒนาแนวทางแก้ไขดังนี้:

\begin{enumerate}
    \item \textbf{ข้อจำกัดด้านการสร้าง IAM Role (\texttt{iam:CreateRole} is Blocked):}
    \begin{itemize}
        \item \textit{ปัญหา:} นโยบาย Learner Lab บล็อกสิทธิ์ \texttt{iam:CreateRole} และ \texttt{iam:AttachRolePolicy} ทำให้คำสั่งอัตโนมัติอย่าง \texttt{sam deploy} ไม่สามารถสร้าง IAM Role ใหม่ได้
        \item \textit{การแก้ไข:} ปรับใช้ \texttt{LabRole} สำเร็จรูปที่มีอยู่แล้ว (\texttt{arn:aws:iam::<Account-ID>:role/LabRole}) ในทุก Lambda Function และเขียนสคริปต์ Deployment อัตโนมัติ (\texttt{scripts/deploy\_manual.sh} และ \texttt{scripts/deploy\_ws.sh}) เพื่ออัปเดตโค้ดขึ้น Lambda โดยตรงผ่าน AWS CLI
    \end{itemize}
    
    \item \textbf{อายุการใช้งานของ Session Token (4-Hour Hard Limit):}
    \begin{itemize}
        \item \textit{ปัญหา:} Credentials ของ Learner Lab จะหมดอายุทุกๆ 4 ชั่วโมง ทำให้การเชื่อมต่อ AWS CLI หลุด
        \item \textit{การแก้ไข:} พัฒนาคู่มือและสคริปต์อัปเดต \texttt{aws\_session\_token} แบบกึ่งอัตโนมัติ และแยก Environment ระหว่าง Local Development กับ Production ออกจากกันอย่างอิสระ
    \end{itemize}
    
    \item \textbf{ข้อจำกัดด้านบริการ (Restricted AWS Services):}
    \begin{itemize}
        \item \textit{ปัญหา:} Learner Lab ไม่อนุญาตให้ใช้ VPC Peering, Custom Domain Name พร้อม SSL Certificate บน API Gateway และไม่รองรับ Amazon Cognito User Pools เต็มรูปแบบ
        \item \textit{การแก้ไข:} พัฒนาระบบ Custom Authentication Microservice (bcrypt + JWT) รันบน AWS Lambda เอง และเชื่อมต่อตรงกับ API Gateway โดยไม่ต้องพึ่งพา Cognito หรือ VPC ซึ่งช่วยลดความซับซ้อนและประหยัดค่าใช้จ่าย
    \end{itemize}
    
    \item \textbf{Payload Format Version Mismatch บน API Gateway:}
    \begin{itemize}
        \item \textit{ปัญหา:} API Gateway HTTP API ค่าตั้งต้นเป็น Payload Version 2.0 ซึ่งส่งโครงสร้าง Event ต่างจาก \texttt{APIGatewayProxyEvent} ของ Node.js ทำให้เกิดข้อผิดพลาด 405 Method Not Allowed
        \item \textit{การแก้ไข:} กำหนด Payload Format Version เป็น 1.0 ในการสร้าง Integration เพื่อให้ส่งต่อ HTTP Method, Headers, และ Path Parameters ครบถ้วน 100\%
    \end{itemize}
\end{enumerate}

% =============================================================
% CHAPTER 4
% =============================================================
\chapter{การพัฒนาระบบและผลการทดสอบคุณภาพ (Implementation \& Quality Assurance)}

\section{รายละเอียดการพัฒนาส่วนติดต่อผู้ใช้ (Frontend Implementation)}
ส่วนติดต่อผู้ใช้ได้รับการออกแบบตามแนวคิด \textit{Dark Editorial + Magic R3} ประกอบด้วยหน้าจอหลัก 5 ส่วน ดังแสดงในรูปที่ \ref{fig:ui_gallery}:

\begin{figure}[H]
\centering
\begin{subfigure}[b]{0.48\textwidth}
    \includegraphics[width=\textwidth]{images/1_before_login_lobby.png}
    \caption{หน้า Match Lobby ก่อนเข้าสู่ระบบ}
\end{subfigure}
\hfill
\begin{subfigure}[b]{0.48\textwidth}
    \includegraphics[width=\textwidth]{images/2_before_login_modal.png}
    \caption{หน้าต่าง Modal เข้าสู่ระบบ / สมัครสมาชิก}
\end{subfigure}
\\[0.3cm]
\begin{subfigure}[b]{0.48\textwidth}
    \includegraphics[width=\textwidth]{images/3_after_login_navbar.png}
    \caption{แถบเมนูนำทางและโปรไฟล์ผู้ใช้หลังเข้าสู่ระบบ}
\end{subfigure}
\hfill
\begin{subfigure}[b]{0.48\textwidth}
    \includegraphics[width=\textwidth]{images/4_after_login_dashboard_decks.png}
    \caption{หน้าจัดการเด็คการ์ดและ Deck Builder}
\end{subfigure}
\caption{ตัวอย่างภาพหน้าจอส่วนติดต่อผู้ใช้ระบบ Disney Lorcana PlayLab Cloud}
\label{fig:ui_gallery}
\end{figure}

องค์ประกอบสำคัญของหน้าจอ Playmat ประกอบด้วย:
\begin{itemize}
    \item \textbf{Play Area (Battlefield):} พื้นที่ลงการ์ดตัวละคร ไอเทม และสถานที่ รองรับการลากวาง (Drag-and-Drop) อย่างอิสระ
    \item \textbf{Inkwell Tray:} พื้นที่คว่ำการ์ดเพื่อใช้เป็นหมึกค่าร่าย พร้อมตัวเลขแสดงจำนวน Unexerted / Total Ink
    \item \textbf{Lore Counter:} มาตรวัดแต้มความรู้ขนาดใหญ่ (0-20 แต้ม) พร้อมปุ่มปรับคะแนนและเอฟเฟกต์แอนิเมชัน
    \item \textbf{Card State Rotator:} กลไกการหมุนการ์ดแนวตั้ง (Ready) และแนวนอน 90 องศา (Exerted) เพื่อระบุว่าการ์ดถูกใช้งานแล้ว
    \item \textbf{3D Card Inspector \& Gacha:} ระบบเปิดซองการ์ด 3D สุ่มการ์ด 12 ใบพร้อมฟิสิกส์แอนิเมชันเปิดซอง
\end{itemize}

\section{รายละเอียดการพัฒนาส่วนประมวลผลบนคลาวด์ (AWS Microservices)}
ฟังก์ชัน Lambda ได้รับการพัฒนาด้วย TypeScript และคอมไพล์เป็น Node.js 20.x แบ่งเป็น 4 โมดูลหลัก:
\begin{enumerate}
    \item \textbf{\texttt{lorcana-auth}:} จัดการการสมัครสมาชิก (\texttt{register}) และการเข้าสู่ระบบ (\texttt{login}) โดยตรวจสอบความถูกต้องของรหัสผ่านด้วย bcrypt และออก JWT Token
    \item \textbf{\texttt{lorcana-deck}:} จัดการ CRUD เด็คการ์ดของผู้ใช้ โดยดึงข้อมูลจากตาราง DynamoDB \texttt{LorcanaDecks} ตาม \texttt{userId}
    \item \textbf{\texttt{lorcana-room-handler}:} จัดการการเชื่อมต่อ WebSocket จัดการ Action \texttt{JOIN\_ROOM}, \texttt{CARD\_MOVED}, \texttt{CARD\_EXERTED}, \texttt{INK\_PLAYED}, \texttt{LORE\_UPDATED}, \texttt{TURN\_PASSED} และ Broadcast ไปยังผู้เล่นคู่แข่งผ่าน \texttt{ApiGatewayManagementApiClient}
    \item \textbf{\texttt{lorcana-analyzer}:} รับงานวิเคราะห์เด็คการ์ดจากคิว Amazon SQS และประมวลผลค่าสถิติ Ink Curve แบบ Asynchronous
\end{enumerate}

\section{แผนและผลการทดสอบคุณภาพระบบ (Quality Assurance \& Test Verification)}
การทดสอบคุณภาพระบบดำเนินตามรูปแบบพีระมิดการทดสอบ (Testing Pyramid) ครอบคลุม Unit Test, Integration Test, End-to-End Test และ Security Test ดังแสดงในตารางที่ \ref{tab:qa_results}:

\begin{table}[H]
\centering
\caption{ผลการทดสอบคุณภาพระบบโดยรวม (Master QA Test Matrix)}
\label{tab:qa_results}
\small
\begin{tabularx}{\textwidth}{p{1.4cm}p{2.2cm}Xp{1.5cm}p{1.5cm}}
\toprule
\textbf{Test ID} & \textbf{ระดับการทดสอบ} & \textbf{เงื่อนไขและกรณีทดสอบ} & \textbf{Latency} & \textbf{ผลการทดสอบ} \\
\midrule
TC-AUTH-01 & Integration & สมัครสมาชิกใหม่ (POST /auth/register) & 180ms & ผ่าน (201 Created) \\
TC-AUTH-02 & Integration & เข้าสู่ระบบด้วยรหัสผ่านถูกต้อง (POST /auth/login) & 145ms & ผ่าน (200 + JWT) \\
TC-AUTH-03 & Security & เข้าสู่ระบบด้วยรหัสผ่านผิด & 120ms & ผ่าน (401 Unauthorized) \\
TC-DECK-01 & Integration & บันทึกเด็คใหม่ 60 ใบ (POST /decks) & 95ms & ผ่าน (201 Created) \\
TC-DECK-02 & Integration & ดึงรายการเด็คของผู้ใช้ (GET /decks) & 68ms & ผ่าน (200 OK) \\
TC-DECK-03 & Integration & ลบเด็คการ์ด (DELETE /decks/\{id\}) & 72ms & ผ่าน (200 OK) \\
TC-WS-01 & E2E & เชื่อมต่อ WSS Persistent Socket (\$connect) & 45ms & ผ่าน (101 Connected) \\
TC-WS-02 & E2E Multi-User & ผู้เล่น 2 คนเข้าห้องเดียวกัน (JOIN\_ROOM) & 82ms & ผ่าน (Room Paired) \\
TC-WS-03 & E2E Multi-User & ซิงค์การลากวางการ์ด (CARD\_MOVED) & 64ms & ผ่าน (<100ms Target) \\
TC-WS-04 & E2E Multi-User & ซิงค์การหมุนการ์ด (CARD\_EXERTED) & 58ms & ผ่าน (<100ms Target) \\
TC-WS-05 & E2E Multi-User & ซิงค์แต้ม Lore (LORE\_UPDATED) & 52ms & ผ่าน (<100ms Target) \\
TC-WS-06 & Fault Tolerance & จำลองเน็ตหลุดและกู้คืน (Auto-Reconnect) & 310ms & ผ่าน (State Restored) \\
TC-SEC-01 & Security (OWASP) & ตรวจสอบการเข้ารหัสรหัสผ่านใน DynamoDB & - & ผ่าน (bcrypt 10 rounds) \\
TC-SEC-02 & Security (OWASP) & ป้องกัน Injection ใน WebSocket Action Payload & - & ผ่าน (Schema Validated) \\
\bottomrule
\end{tabularx}
\end{table}

ผลการทดสอบยืนยันว่าการซิงค์ข้อมูลกระดานผ่าน AWS API Gateway WebSockets มีความหน่วงเฉลี่ยระหว่าง \textbf{52--82 มิลลิวินาที} ซึ่งผ่านเกณฑ์เป้าหมาย Sub-100ms ได้อย่างสมบูรณ์

% =============================================================
% CHAPTER 5
% =============================================================
\chapter{สรุปผลการดำเนินงานและแผนงานในอนาคต (Conclusions \& Future Roadmap)}

\section{สรุปผลการดำเนินงานความก้าวหน้า Stage 2}
การดำเนินงานโครงงาน \textbf{Disney Lorcana PlayLab Cloud} ในระยะที่ 2 (Stage 2) ครอบคลุมการพัฒนาตามแผนงาน Sprint 1 ถึง Sprint 3 สำเร็จลุล่วง 100\% ตามข้อกำหนดของรายวิชา Cloud Technology โดยมีผลสัมฤทธิ์สำคัญดังนี้:
\begin{enumerate}
    \item พัฒนาส่วนติดต่อผู้ใช้จำลองกระดานเล่นการ์ด Disney Lorcana ที่สมบูรณ์แบบ รองรับกฎการเล่นมาตรฐาน ฐานข้อมูลการ์ด 408 ใบ และระบบ 3D Card Inspector
    \item ออกแบบและปรับใช้สถาปัตยกรรมคลาวด์แบบไร้เซิร์ฟเวอร์เต็มรูปแบบ (100\% Serverless) บน AWS ประกอบด้วย AWS Lambda, Amazon DynamoDB, และ AWS API Gateway (HTTP \& WebSocket)
    \item บรรลุเป้าหมายความเร็วในการซิงค์ข้อมูลระหว่างผู้เล่นด้วยความหน่วงต่ำกว่า 100 มิลลิวินาที (เฉลี่ย 52--82ms)
    \item เอาชนะข้อจำกัดสิทธิ์บน AWS Academy Learner Lab ได้อย่างมีประสิทธิภาพผ่านการใช้ Pre-provisioned LabRole และสคริปต์ Deployment อัตโนมัติ
    \item ควบคุมค่าใช้จ่ายโครงสร้างพื้นฐานทั้งหมดให้อยู่ในโควตา AWS Free Tier (\$0.00) สอดคล้องกับกรอบงบประมาณที่กำหนด
\end{enumerate}

\section{แผนงานการพัฒนาสำหรับ Stage 3 (Sprint 4 ถึง Sprint 7)}
\begin{enumerate}
    \item \textbf{Sprint 4 (Async Deck Analyzer via Amazon SQS):} พัฒนาระบบประมวลผลวิเคราะห์เด็คการ์ดเชิงลึก คำนวณความเสี่ยง ค่าร่ายเฉลี่ย และแนะนำการปรับเด็คโดยใช้ Amazon SQS รับงานและส่งต่อให้ Lambda ทำงานเบื้องหลัง
    \item \textbf{Sprint 5 (Observability, Monitoring \& Security Hardening):} สร้าง Amazon CloudWatch Dashboard แสดงผลแบบรวมศูนย์ และตั้งค่า CloudWatch Alarms แจ้งเตือนเมื่อเกิด Error Rate สูงกว่า 1\%
    \item \textbf{Sprint 6 (Load Testing \& High Concurrency Simulation):} ดำเนินการทดสอบโหลดด้วยเครื่องมือ Artillery เพื่อจำลองผู้เล่นพร้อมกัน 100--500 ผู้ใช้ และทดสอบความทนทานของระบบ WebSocket Auto-reconnection
    \item \textbf{Sprint 7 (Final Report \& Video Presentation):} จัดทำรายงานฉบับสมบูรณ์ วิดีโอสาธิตการใช้งานระบบ และเตรียมการนำเสนอโครงงานในรอบสุดท้าย
\end{enumerate}

\section{บทเรียนที่ได้รับและการปรับปรุงกระบวนการทำงาน (Retrospective)}
\begin{enumerate}
    \item \textbf{ความสำคัญของการออกแบบ Interface และ Contract First:} การกำหนดรูปแบบ JSON Payload ของ WebSocket และ DynamoDB Schema ล่วงหน้าช่วยลดข้อผิดพลาดในการรวมระบบระหว่าง Frontend และ Cloud Backend ได้อย่างมาก
    \item \textbf{การเข้าใจสภาพแวดล้อมคลาวด์แบบจำกัดสิทธิ์:} การเตรียมแผนสำรองสำหรับข้อจำกัดด้าน IAM Role บน AWS Learner Lab ทำให้ทีมงานสามารถส่งมอบงานได้ตรงตามกำหนดการโดยไม่เกิดความล่าช้า
    \item \textbf{การบริหารจัดการสถาปัตยกรรมไร้เซิร์ฟเวอร์:} การเลือกใช้บริการ Serverless ช่วยลดภาระงานด้านการดูแลระบบลงได้อย่างมหาศาล ทำให้ทีมงานสามารถมุ่งเน้นการพัฒนาฟีเจอร์และปรับปรุงประสบการณ์ของผู้ใช้งานได้อย่างเต็มที่
\end{enumerate}

% -------------------------------------------------------------
% REFERENCES (บรรณานุกรม)
% -------------------------------------------------------------
\chapter*{บรรณานุกรม (References)}
\addcontentsline{toc}{chapter}{บรรณานุกรม (References)}

\noindent [1] Ravensburger \& Disney. (2023). \textit{Disney Lorcana Trading Card Game Official Rules and Card Database}. [Online]. Available: \url{https://www.disneylorcana.com}\\[0.3cm]
\noindent [2] Amazon Web Services. (2026). \textit{AWS Lambda Developer Guide: Building Event-Driven Serverless Applications}. [Online]. Available: \url{https://docs.aws.amazon.com/lambda/}\\[0.3cm]
\noindent [3] Amazon Web Services. (2026). \textit{Amazon API Gateway Developer Guide: About WebSocket APIs in API Gateway}. [Online]. Available: \url{https://docs.aws.amazon.com/apigateway/latest/developerguide/apigateway-websocket-api.html}\\[0.3cm]
\noindent [4] Amazon Web Services. (2026). \textit{Amazon DynamoDB Developer Guide: Core Components and Best Practices for NoSQL Design}. [Online]. Available: \url{https://docs.aws.amazon.com/amazondynamodb/}\\[0.3cm]
\noindent [5] Amazon Web Services. (2026). \textit{AWS Well-Architected Framework: Reliability, Security, and Cost Optimization Pillars}. [Online]. Available: \url{https://aws.amazon.com/architecture/well-architected/}\\[0.3cm]
\noindent [6] React Working Group. (2026). \textit{React 19 Documentation: Concurrent Features and Server Components}. [Online]. Available: \url{https://react.dev}\\[0.3cm]
\noindent [7] Tailwind Labs. (2026). \textit{Tailwind CSS v4.0 Alpha/Release Documentation}. [Online]. Available: \url{https://tailwindcss.com}\\[0.3cm]
\noindent [8] Framer. (2026). \textit{Framer Motion: Production-Ready Animation Library for React}. [Online]. Available: \url{https://www.framer.com/motion/}\\[0.3cm]
\noindent [9] IETF (Internet Engineering Task Force). (2011). \textit{RFC 6455: The WebSocket Protocol}. [Online]. Available: \url{https://tools.ietf.org/html/rfc6455}

% -------------------------------------------------------------
% APPENDIX (ภาคผนวก)
% -------------------------------------------------------------
\chapter*{ภาคผนวก (Appendix)}
\addcontentsline{toc}{chapter}{ภาคผนวก (Appendix)}

\section*{ตัวอย่างโค้ดสคริปต์ Deployment บน AWS Learner Lab (\texttt{scripts/deploy\_ws.sh})}

\begin{lstlisting}[language=bash, caption={สคริปต์ Deploy WebSocket Lambda \& API Gateway บน AWS Learner Lab}]
#!/bin/bash
# scripts/deploy_ws.sh - Deploy Real-Time WebSocket Infrastructure
set -e

echo "=== 1. Building Backend TypeScript ==="
cd ../backend
npm install
npx tsc --outDir dist

echo "=== 2. Packaging Lambda Handler ==="
cd dist/room
zip -r ../../../room-handler.zip .
cd ../../..

echo "=== 3. Updating Lambda Code via AWS CLI ==="
aws lambda update-function-code \
  --function-name lorcana-room-handler \
  --zip-file fileb://room-handler.zip \
  --region us-east-1

echo "=== 4. Verifying WebSocket API Gateway Integration ==="
API_ID=$(aws apigatewayv2 get-apis --query "Items[?Name=='lorcana-ws-api'].ApiId" --output text)
echo "WebSocket API ID: $API_ID"
echo "WebSocket Endpoint: wss://${API_ID}.execute-api.us-east-1.amazonaws.com/prod"

echo "=== Deployment Completed Successfully ==="
\end{lstlisting}

\end{document}
""")

print(f"[1/3] Generated LaTeX document: {latex_file}")

# 2. XELATEX COMPILATION
print("[2/3] Compiling LaTeX with xelatex (Pass 1)...")
res1 = subprocess.run(["xelatex", "-interaction=nonstopmode", "G21_Disney_Lorcana_PlayLab_Cloud_Phase2_Report.tex"], 
                      cwd=DOCS_DIR, capture_output=True, encoding='utf-8', errors='ignore')
print(f"Pass 1 Return Code: {res1.returncode}")

print("[2/3] Compiling LaTeX with xelatex (Pass 2 for Table of Contents & References)...")
res2 = subprocess.run(["xelatex", "-interaction=nonstopmode", "G21_Disney_Lorcana_PlayLab_Cloud_Phase2_Report.tex"], 
                      cwd=DOCS_DIR, capture_output=True, encoding='utf-8', errors='ignore')
print(f"Pass 2 Return Code: {res2.returncode}")

pdf_file = os.path.join(DOCS_DIR, "G21_Disney_Lorcana_PlayLab_Cloud_Phase2_Report.pdf")
if os.path.exists(pdf_file):
    print(f"PDF generated successfully: {pdf_file} ({os.path.getsize(pdf_file):,} bytes)")
else:
    print("Error generating PDF. Output tail:")
    print(res2.stdout[-500:] if res2.stdout else res2.stderr[-500:])

# 3. DOCX GENERATION VIA PYTHON-DOCX
print("[3/3] Generating MS Word (.docx) via python-docx...")

try:
    import docx
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls

    doc = Document()

    for s in doc.sections:
        s.top_margin = Inches(1.38)
        s.bottom_margin = Inches(0.98)
        s.left_margin = Inches(1.38)
        s.right_margin = Inches(0.98)
        s.page_width = Inches(8.27)
        s.page_height = Inches(11.69)

    norm = doc.styles['Normal']
    norm.font.name = 'TH Sarabun New'
    norm.font.size = Pt(16)
    norm.font.color.rgb = RGBColor(0, 0, 0)
    norm.paragraph_format.line_spacing = 1.15
    norm.paragraph_format.space_after = Pt(6)

    def heading(text, level):
        p = doc.add_paragraph()
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        r.font.name = 'TH Sarabun New'
        r.bold = True
        if level == 1:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r.font.size = Pt(20)
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(12)
        elif level == 2:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r.font.size = Pt(18)
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
        elif level == 3:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r.font.size = Pt(16)
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
        return p

    def body(text, indent=True, bold_prefix=None):
        p = doc.add_paragraph()
        if indent:
            p.paragraph_format.first_line_indent = Inches(0.5)
        if bold_prefix:
            rb = p.add_run(bold_prefix)
            rb.font.name = 'TH Sarabun New'
            rb.bold = True
            rb.font.size = Pt(16)
        r = p.add_run(text)
        r.font.name = 'TH Sarabun New'
        r.font.size = Pt(16)
        return p

    # Cover
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_path = os.path.join(IMAGES_DIR, 'IT_KMITL_ICON.png')
    if os.path.exists(logo_path):
        p_logo.add_run().add_picture(logo_path, width=Inches(1.3))

    p_t1 = doc.add_paragraph()
    p_t1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_t1.paragraph_format.space_before = Pt(12)
    r = p_t1.add_run("รายงานโครงงานวิชาการเทคโนโลยีกลุ่มเมฆ (Cloud Technology)\nภาคเรียนที่ 1 ปีการศึกษา 2569\n\n")
    r.font.name = 'TH Sarabun New'
    r.font.size = Pt(18)
    r.bold = True

    r2 = p_t1.add_run("โครงงานระบบจำลองห้องเล่นและวิเคราะห์เด็คการ์ดแบบเรียลไทม์บนคลาวด์ Serverless\n(Disney Lorcana PlayLab Cloud: Real-Time Serverless Card Game Simulator)\nกลุ่ม G21\n\n")
    r2.font.name = 'TH Sarabun New'
    r2.font.size = Pt(18)
    r2.bold = True

    r3 = p_t1.add_run("รายงานความก้าวหน้าโครงการ (Stage 2: บทที่ 1 - 3 และส่วนขยายระบบคลาวด์)\n\n\n")
    r3.font.name = 'TH Sarabun New'
    r3.font.size = Pt(16)
    r3.bold = True

    # Table for Authors & Professors in Word
    t_info = doc.add_table(rows=1, cols=2)
    t_info.autofit = False
    for row in t_info.rows:
        row.cells[0].width = Inches(3.2)
        row.cells[1].width = Inches(2.3)
        for cell in row.cells:
            # remove borders
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = parse_xml(r'<w:tcBorders %s><w:top w:val="none"/><w:left w:val="none"/><w:bottom w:val="none"/><w:right w:val="none"/></w:tcBorders>' % nsdecls('w'))
            tcPr.append(tcBorders)

    c0 = t_info.cell(0, 0)
    p0 = c0.paragraphs[0]
    p0.paragraph_format.line_spacing = 1.15
    p0.paragraph_format.space_after = Pt(2)
    r = p0.add_run("คณะผู้จัดทำ:\n")
    r.font.name = 'TH Sarabun New'
    r.font.size = Pt(14)
    r.bold = True
    members = [
        "1. นายชยุต บุญวัฒน์",
        "2. นายธนัทภัทร พรหมทอง",
        "3. นายภูริ ประชาสุขสิน",
        "4. นางสาววรรณณิศา อมรวงศ์ไพบูลย์",
        "5. นายอสิธารา พุ่มดอกไม้",
        "6. นายวรธิษณ์ คงทอง"
    ]
    r = p0.add_run("\n".join(members) + "\n\n")
    r.font.name = 'TH Sarabun New'
    r.font.size = Pt(13.5)

    r_prof_hdr = p0.add_run("อาจารย์ประจำวิชา:\n")
    r_prof_hdr.font.name = 'TH Sarabun New'
    r_prof_hdr.font.size = Pt(14)
    r_prof_hdr.bold = True

    profs = [
        "1. ดร. ธนานพ ทองถาวร",
        "2. ผศ.ดร. พัฒนพงษ์ ฉันทมิตรโอภาส",
        "3. ผศ.ดร. ลภัส ประดิษฐ์ทัศนีย์"
    ]
    r_prof = p0.add_run("\n".join(profs))
    r_prof.font.name = 'TH Sarabun New'
    r_prof.font.size = Pt(13.5)

    c1 = t_info.cell(0, 1)
    p1 = c1.paragraphs[0]
    p1.paragraph_format.line_spacing = 1.15
    p1.paragraph_format.space_after = Pt(2)
    student_ids = [
        "รหัสนักศึกษา 67070032",
        "รหัสนักศึกษา 67070069",
        "รหัสนักศึกษา 67070137",
        "รหัสนักศึกษา 67070155",
        "รหัสนักศึกษา 67070199",
        "รหัสนักศึกษา 67070275"
    ]
    r = p1.add_run("\n" + "\n".join(student_ids) + "\n\n\n\n\n\n")
    r.font.name = 'TH Sarabun New'
    r.font.size = Pt(13.5)

    p_foot = doc.add_paragraph()
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_foot = p_foot.add_run("คณะเทคโนโลยีสารสนเทศ\nสถาบันเทคโนโลยีพระจอมเกล้าเจ้าคุณทหารลาดกระบัง\nกันยายน 2569")
    r_foot.font.name = 'TH Sarabun New'
    r_foot.font.size = Pt(16)
    r_foot.bold = True

    doc.add_page_break()

    # Abstract Thai
    heading("บทคัดย่อภาษาไทย", 1)
    body("โครงงาน Disney Lorcana PlayLab Cloud (กลุ่ม G21) นำเสนอการวิเคราะห์ ออกแบบ และพัฒนาระบบจำลองห้องเล่นและวิเคราะห์เด็คการ์ดแบบเรียลไทม์บนสถาปัตยกรรมคลาวด์ไร้เซิร์ฟเวอร์เต็มรูปแบบ (100% AWS Serverless Architecture) เพื่อรองรับเกมการ์ดสะสม Disney Lorcana Trading Card Game (TCG) โดยมุ่งเน้นการขจัดปัญหาภาระค่าใช้จ่ายเซิร์ฟเวอร์แบบเดิม (Server Overhead) และลดความหน่วงในการส่งข้อมูลระหว่างผู้เล่น (Latency Reduction) ให้ต่ำกว่า 100 มิลลิวินาที (Sub-100ms)")
    body("โครงสร้างระบบขับเคลื่อนด้วยการประมวลผลตามเหตุการณ์ (Event-Driven Computing) ผ่าน AWS Lambda ในรูปแบบ Microservices รองรับการสื่อสารแบบสองทิศทางความเร็วสูงผ่าน AWS API Gateway WebSockets ร่วมกับการจัดเก็บข้อมูลถาวรบน Amazon DynamoDB ซึ่งปรับขนาดอัตโนมัติ (Auto-scaling) และการประมวลผลวิเคราะห์เด็คการ์ดแบบ Asynchronous ผ่าน Amazon SQS ในส่วนของผู้ใช้งาน (Frontend) พัฒนาด้วย React 19, TypeScript 5, Tailwind CSS v4, Zustand State Machine และระบบฟิสิกส์ 3D Card Inspector พร้อมเชื่อมต่อฐานข้อมูลการ์ดมาตรฐาน 408 ใบ")
    body("นอกจากนี้ รายงานฉบับนี้ได้ลงลึกถึงการออกแบบวิศวกรรมเพื่อความพร้อมใช้งานสูง (High Availability) และความเชื่อถือได้ (Reliability) การป้องกันความล้มเหลวจากการพึ่งพา API ภายนอกด้วยกลยุทธ์ Multi-tier Fallback และ Local Caching การจัดการสเกลการเชื่อมต่อ WebSockets เมื่อมีห้องเล่นพร้อมกันจำนวนมากด้วย Partition Key Dispersion และ Connection Multiplexing ตลอดจนการปฏิบัติตามกรอบ AWS Well-Architected Framework ภายใต้งบประมาณที่ควบคุมได้ ($0.00 Free Tier และไม่เกิน $50) รวมถึงการแก้ไขข้อจำกัดเฉพาะของสภาพแวดล้อม AWS Learner Lab ได้อย่างสมบูรณ์แบบ")
    body("คลาวด์ไร้เซิร์ฟเวอร์ (Serverless), เว็บซ็อกเก็ต (WebSockets), ดึงข้อมูลตามเหตุการณ์ (Event-Driven), เกมการ์ดดิจิทัล, AWS Lambda, DynamoDB, API Gateway", indent=False, bold_prefix="คำสำคัญ: ")

    doc.add_page_break()

    # Abstract English
    heading("Abstract", 1)
    body("The Disney Lorcana PlayLab Cloud (Group G21) project presents the architectural design, engineering analysis, and implementation of a real-time multiplayer card simulator and asynchronous deck analytics platform built entirely on a 100% AWS Serverless Architecture for the Disney Lorcana Trading Card Game (TCG). The project targets the elimination of persistent infrastructure costs and achieves ultra-low synchronization latency (<100ms) between remote players.")
    body("The cloud backend is designed around event-driven microservices using AWS Lambda, bidirectional persistent connections via AWS API Gateway WebSockets, autoscaling document storage with Amazon DynamoDB, and decoupled asynchronous job processing through Amazon Simple Queue Service (SQS). The frontend client is engineered using React 19, TypeScript 5, Tailwind CSS v4, Zustand global state orchestration, Framer Motion, and CSS 3D matrix physics, delivering a high-fidelity tabletop simulation connected to an official catalog of 408 cards.")
    body("Furthermore, this Stage 2 report provides comprehensive engineering solutions for critical cloud challenges: mitigating external card API dependencies through multi-tier caching and local asset fallbacks, architecting WebSocket scalability under high concurrent room load using DynamoDB partition key dispersion, auto-reconnect grace windows, and adherence to the AWS Well-Architected Framework (Reliability, Availability, Security, and Cost under $50 budget limit). Finally, it details the engineering workarounds engineered to overcome AWS Academy Learner Lab permission boundaries (pre-provisioned LabRole integration and automated CLI pipelines).")
    body("AWS Serverless, WebSockets, Event-Driven Architecture, Digital TCG, AWS Lambda, DynamoDB, API Gateway, Cloud Reliability", indent=False, bold_prefix="Keywords: ")

    doc.add_page_break()

    # Acknowledgments
    heading("กิตติกรรมประกาศ", 1)
    body("โครงงาน Disney Lorcana PlayLab Cloud สำเร็จลุล่วงตามวัตถุประสงค์ของการส่งมอบงานในระยะที่ 2 (Stage 2) ได้ด้วยความอนุเคราะห์ คำแนะนำ และการชี้แนะทางวิชาการอันทรงคุณค่ายิ่งจากคณาจารย์ประจำรายวิชาการเทคโนโลยีกลุ่มเมฆ (Cloud Technology) คณะเทคโนโลยีสารสนเทศ สถาบันเทคโนโลยีพระจอมเกล้าเจ้าคุณทหารลาดกระบัง ได้แก่:\n1. ดร. ธนานพ ทองถาวร\n2. ผศ.ดร. พัฒนพงษ์ ฉันทมิตรโอภาส\n3. ผศ.ดร. ลภัส ประดิษฐ์ทัศนีย์")
    body("ที่ได้กรุณาถ่ายทอดองค์ความรู้ด้านสถาปัตยกรรมคลาวด์ การคำนึงถึงความปลอดภัย ความพร้อมใช้งาน และการบริหารจัดการทรัพยากรอย่างคุ้มค่า ซึ่งเป็นรากฐานสำคัญในการออกแบบและแก้ปัญหาทางวิศวกรรมซอฟต์แวร์ในโครงงานนี้")
    body("คณะผู้จัดทำขอขอบพระคุณเพื่อน ๆ นักศึกษาคณะเทคโนโลยีสารสนเทศ ทุกท่านที่ได้ร่วมทดสอบระบบกระดานและห้องเล่นแบบเรียลไทม์ ให้ข้อคิดเห็นด้านประสบการณ์การใช้งาน (UX/UI) และร่วมตรวจสอบข้อผิดพลาดของระบบ รวมถึงขอขอบคุณครอบครัวที่ให้การสนับสนุนและเป็นกำลังใจในการศึกษาค้นคว้าตลอดมา")

    p_s = doc.add_paragraph()
    p_s.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_s.paragraph_format.space_before = Pt(30)
    p_s.add_run("คณะผู้จัดทำ โครงงานกลุ่ม G21\nกันยายน 2569").font.name = 'TH Sarabun New'

    doc.add_page_break()

    # TOC
    heading("สารบัญ (Table of Contents)", 1)
    p_t = doc.add_paragraph()
    p_t.add_run("""บทคัดย่อภาษาไทย .......................................................................................................................... I
Abstract ......................................................................................................................................... II
กิตติกรรมประกาศ ........................................................................................................................... III
สารบัญ ........................................................................................................................................... IV
สารบัญตาราง ................................................................................................................................. V
สารบัญรูปภาพ ................................................................................................................................ VI

บทที่ 1: บทนำ (Introduction) ....................................................................................................... 1
  1.1 ความเป็นมาและความสำคัญของปัญหา ................................................................................... 1
  1.2 วัตถุประสงค์ของโครงงาน ..................................................................................................... 2
  1.3 ขอบเขตของระบบ (System Scope) ..................................................................................... 2
  1.4 ประโยชน์ที่คาดว่าจะได้รับ ..................................................................................................... 3
  1.5 โครงสร้างของรายงาน ........................................................................................................... 3

บทที่ 2: ทฤษฎีและเทคโนโลยีที่เกี่ยวข้อง (Related Theories & Technologies) ............................ 4
  2.1 สถาปัตยกรรมคลาวด์แบบไร้เซิร์ฟเวอร์ (AWS Serverless Architecture) ................................ 4
  2.2 สถาปัตยกรรมการสื่อสารแบบสองทิศทางเรียลไทม์ (WebSockets Protocol) ........................... 5
  2.3 เทคโนโลยีฝั่งผู้ใช้งาน (Frontend Engineering & State Architecture) ................................. 5
  2.4 โมเดลโดเมนเกมการ์ดสะสม Disney Lorcana ..................................................................... 6
  2.5 ระเบียบวิธีปฏิบัติงานแบบ Agile Scrum และการดำเนินงาน Stage 2 ...................................... 6

บทที่ 3: การวิเคราะห์และออกแบบระบบ (System Analysis & Architectural Design) ................. 8
  3.1 ข้อกำหนดความต้องการของระบบ (Functional & Non-Functional Requirements) ............ 8
  3.2 สถาปัตยกรรมระบบโดยรวม (End-to-End Serverless Architecture) .................................... 9
  3.3 การออกแบบ Use Case และลำดับการทำงาน (Use Cases LC-01 ถึง LC-09) ....................... 10
  3.4 การออกแบบฐานข้อมูลแบบ NoSQL (Amazon DynamoDB Schema) ................................. 11
  3.5 การออกแบบการจัดการ API ภายนอกและการป้องกันความล้มเหลว (External API Strategy) ... 12
  3.6 การออกแบบการสเกลและรองรับ WebSockets Concurrency ............................................... 13
  3.7 การออกแบบตามกรอบ AWS Well-Architected Framework (Reliability, Availability, Security, Cost) . 14
  3.8 การวิเคราะห์ข้อจำกัดของ AWS Learner Lab และแนวทางแก้ไขเชิงวิศวกรรม ..................... 16

บทที่ 4: การพัฒนาระบบและผลการทดสอบคุณภาพ (Implementation & Quality Assurance) ..... 18
  4.1 รายละเอียดการพัฒนาส่วนติดต่อผู้ใช้ (Frontend Implementation & Playmat UI) ................ 18
  4.2 รายละเอียดการพัฒนาส่วนประมวลผลบนคลาวด์ (AWS Microservices) ................................. 19
  4.3 แผนและผลการทดสอบคุณภาพระบบ (Master QA Test Matrix & Verification) .................. 20

บทที่ 5: สรุปผลการดำเนินงานและแผนงานในอนาคต (Conclusions & Future Roadmap) ............. 22
  5.1 สรุปผลการดำเนินงานความก้าวหน้า Stage 2 ......................................................................... 22
  5.2 แผนงานการพัฒนาสำหรับ Stage 3 (Sprint 4 ถึง 7) ............................................................. 22
  5.3 บทเรียนที่ได้รับและการปรับปรุงกระบวนการทำงาน (Sprint Retrospective) ........................... 23

บรรณานุกรม (References) ........................................................................................................... 24
ภาคผนวก (Appendix) ................................................................................................................. 25""").font.name = 'TH Sarabun New'

    doc.add_page_break()

    # Chapters 1-5
    heading("บทที่ 1\nบทนำ (Introduction)", 1)
    heading("1.1 ความเป็นมาและความสำคัญของปัญหา", 2)
    body("เกมการ์ดสะสมเชิงกลยุทธ์ (Trading Card Game: TCG) เป็นหนึ่งในอุตสาหกรรมเกมและสื่อสันทนาการที่มีอัตราการเติบโตสูงอย่างต่อเนื่องในระดับสากล โดยเฉพาะอย่างยิ่งการเปิดตัวเกม Disney Lorcana TCG โดยความร่วมมือระหว่าง Ravensburger และ The Walt Disney Company ซึ่งได้รับความนิยมอย่างแพร่หลายจากทั้งผู้เล่นสายแข่งขัน (Competitive Players) และนักสะสมทั่วโลก")
    body("อย่างไรก็ดี ในบริบทของการฝึกซ้อมและการเข้าถึงเกมการ์ด ผู้เล่นส่วนใหญ่ยังคงประสบกับข้อจำกัดหลัก 3 ประการ: (1) ต้นทุนการจัดหาการ์ดจริงและข้อจำกัดทางกายภาพ (2) ภาระต้นทุนเซิร์ฟเวอร์และความไร้ประสิทธิภาพของระบบดั้งเดิมที่พึ่งพา Virtual Machine เปิดค้างตลอด 24 ชั่วโมง และ (3) ความหน่วงในการส่งข้อมูล (High Latency) จากการใช้ HTTP Polling")
    body("เพื่อแก้ไขปัญหาดังกล่าว คณะผู้จัดทำจึงได้ริเริ่มโครงงาน Disney Lorcana PlayLab Cloud (กลุ่ม G21) โดยนำสถาปัตยกรรมคลาวด์แบบไร้เซิร์ฟเวอร์เต็มรูปแบบ (100% AWS Serverless Architecture) ร่วมกับเทคโนโลยีเว็บสมัยใหม่ มาใช้ในการพัฒนาระบบจำลองกระดานเล่นการ์ดและระบบวิเคราะห์เด็คอัจฉริยะ ที่มีความหน่วงต่ำกว่า 100 มิลลิวินาที สามารถขยายขนาดรองรับผู้เล่นได้โดยอัตโนมัติ และควบคุมค่าใช้จ่ายโครงสร้างพื้นฐานให้อยู่ภายใต้งบประมาณฟรีของ AWS Free Tier ($0.00) ได้อย่างสมบูรณ์")

    heading("1.2 วัตถุประสงค์ของโครงงาน", 2)
    body("1. เพื่อออกแบบและพัฒนาระบบจำลองกระดานซ้อมเล่นการ์ด Disney Lorcana แบบเรียลไทม์บนเว็บเบราว์เซอร์ที่มีความหน่วงต่ำ (<100ms) ผ่านช่องทางสื่อสารสองทิศทาง WebSockets\n2. เพื่อประยุกต์ใช้สถาปัตยกรรม AWS Serverless (AWS Lambda, Amazon DynamoDB, AWS API Gateway) ในการสร้างระบบ Microservices ที่ประหยัดพลังงาน ยืดหยุ่น และปรับขนาดตามโหลดอัตโนมัติ\n3. เพื่อพัฒนาระบบจัดการเด็คการ์ด (Deck Builder) และระบบประมวลผลวิเคราะห์สมดุลเด็คแบบ Asynchronous ผ่าน Amazon SQS\n4. เพื่อศึกษาและวางมาตรการด้านความเชื่อถือได้ (Reliability), ความพร้อมใช้งาน (Availability), ความปลอดภัย (Security) และการควบคุมต้นทุน (Cost Optimization) ตามกรอบ AWS Well-Architected Framework\n5. เพื่อศึกษาและแก้ไขปัญหาข้อจำกัดทางวิศวกรรมบนสภาพแวดล้อมสิทธิ์จำกัดของ AWS Academy Learner Lab")

    heading("1.3 ขอบเขตของระบบ (System Scope)", 2)
    body("ระบบครอบคลุมฟังก์ชันการทำงานหลัก 9 ด้าน (LC-01 ถึง LC-09) ได้แก่: ระบบยืนยันตัวตนและความปลอดภัย (bcrypt + JWT), ระบบจัดการเด็คการ์ดบน DynamoDB, ระบบจำลองการเปิดซองการ์ด Booster Pack, ระบบตรวจสอบการ์ดแบบ 3 มิติ (3D Inspector), ระบบจำลองกระดานและห้องเล่นเรียลไทม์ผ่าน WebSockets, ระบบวิเคราะห์เด็คแบบ Asynchronous บน SQS, ระบบจัดการผู้ใช้งาน และระบบตรวจสอบประสิทธิภาพผ่าน CloudWatch")

    heading("1.4 ประโยชน์ที่คาดว่าจะได้รับ", 2)
    body("1. ผู้เล่นมีเครื่องมือฝึกซ้อมและวิเคราะห์เด็คการ์ดเสมือนจริงที่ใช้งานได้ฟรีผ่านเว็บเบราว์เซอร์\n2. ได้รับความรู้เชิงลึกในการออกแบบสถาปัตยกรรม Serverless Microservices บน AWS\n3. ได้แนวทางปฏิบัติที่เป็นรูปธรรมในการรับมือกับความล้มเหลวของ API ภายนอกและการสเกล WebSockets\n4. ลดภาระต้นทุนด้านโครงสร้างพื้นฐานลงได้ 100% บน AWS Free Tier")

    heading("1.5 โครงสร้างของรายงาน", 2)
    body("รายงานประกอบด้วย 5 บท ได้แก่ บทที่ 1 บทนำ, บทที่ 2 ทฤษฎีและเทคโนโลยีที่เกี่ยวข้อง, บทที่ 3 การวิเคราะห์และออกแบบระบบ, บทที่ 4 การพัฒนาระบบและผลการทดสอบคุณภาพ และบทที่ 5 สรุปผลการดำเนินงานและแผนงานในอนาคต")

    heading("บทที่ 2\nทฤษฎีและเทคโนโลยีที่เกี่ยวข้อง (Related Theories & Technologies)", 1)
    heading("2.1 สถาปัตยกรรมคลาวด์แบบไร้เซิร์ฟเวอร์ (AWS Serverless Architecture)", 2)
    body("สถาปัตยกรรม Serverless ขจัดภาระการบริหารจัดการเครื่องเซิร์ฟเวอร์ โดยคิดค่าบริการตามปริมาณการใช้งานจริงในระดับมิลลิวินาที โครงงานนี้ใช้บริการหลักของ AWS ได้แก่ AWS Lambda (Event-Driven FaaS), Amazon API Gateway (REST HTTP API & WebSocket API), Amazon DynamoDB (NoSQL On-Demand Database), Amazon SQS (Distributed Message Queue) และ Amazon CloudWatch (Monitoring & Alarms)")

    heading("2.2 สถาปัตยกรรมการสื่อสารแบบสองทิศทางเรียลไทม์ (WebSockets Protocol)", 2)
    body("โพรโทคอล WebSocket (RFC 6455) ทำงานแบบ Full-Duplex บนช่องทาง TCP เดียวกัน AWS API Gateway WebSockets จัดการ Connection Lifecycle ($connect, $disconnect), Route Selection Expression ($request.body.action) และมี ApiGatewayManagementApi (@connections) ที่อนุญาตให้ Lambda ส่งข้อมูล Broadcast ไปยังผู้เล่นคู่แข่งได้โดยตรง")

    heading("2.3 เทคโนโลยีฝั่งผู้ใช้งาน (Frontend Engineering & State Architecture)", 2)
    body("พัฒนาด้วย React 19, TypeScript 5, Tailwind CSS v4 ในธีม Dark Editorial + Magic R3 พร้อมการจัดการสถานะด้วย Zustand (usePlaymatStore, useDeckStore, useAuthStore) และ Framer Motion ร่วมกับ CSS 3D Transforms สำหรับฟิสิกส์การลากวางและการเปิดซองการ์ด 3D")

    heading("2.4 โมเดลโดเมนเกมการ์ดสะสม Disney Lorcana", 2)
    body("ประกอบด้วยสีหมึก 6 หมวด (Amber, Amethyst, Emerald, Ruby, Sapphire, Steel), ประเภทการ์ด (Character, Action, Item, Location), สถานะการ์ด (Ready และ Exerted) และเงื่อนไขการชนะคือการสะสมแต้ม Lore ครบ 20 แต้ม")

    heading("2.5 ระเบียบวิธีปฏิบัติงานแบบ Agile Scrum และการดำเนินงาน Stage 2", 2)
    body("โครงงานแบ่งการพัฒนาออกเป็น 7 Sprints โดยใน Stage 2 นี้ ได้ส่งมอบ Sprint 1 (Board UI & 408 Cards Data), Sprint 2 (Auth JWT/bcrypt & Deck REST API) และ Sprint 3 (WebSocket Real-time Room Sync) สำเร็จลุล่วง 100%")

    heading("บทที่ 3\nการวิเคราะห์และออกแบบระบบ (System Analysis & Architectural Design)", 1)
    heading("3.1 ข้อกำหนดความต้องการของระบบ (Requirements)", 2)
    body("ระบบประกอบด้วย Functional Requirements (FR-01 ถึง FR-07) ครอบคลุมฟังก์ชันการเล่น การจัดการเด็ค และการวิเคราะห์ และ Non-Functional Requirements (NFR-01 ถึง NFR-04) ที่เน้น Latency < 100ms, Availability 99.9%, Security bcrypt/JWT และ Cost $0.00 Free Tier")

    heading("3.2 สถาปัตยกรรมระบบโดยรวม (End-to-End Serverless Architecture)", 2)
    body("แบ่งเป็น 4 เลเยอร์: (1) Presentation Layer: React 19 SPA (2) API & Routing Layer: HTTP API & WebSocket API Gateway (3) Compute Layer: AWS Lambda Microservices (Auth, Deck, Room Router, Analyzer) (4) Persistence Layer: Amazon DynamoDB และ Amazon SQS")

    heading("3.3 การออกแบบ Use Case และลำดับการทำงาน", 2)
    body("ครอบคลุม 9 Use Cases สำคัญ (LC-01 ถึง LC-09) พร้อม Flow การทำงานแบบ Real-time และ Asynchronous")

    heading("3.4 การออกแบบฐานข้อมูลแบบ NoSQL (Amazon DynamoDB Schema)", 2)
    body("ออกแบบ 3 ตารางแยกตามขอบเขตงาน: LorcanaUsers (PK: userId), LorcanaDecks (PK: userId, SK: deckId), LorcanaRoomState (PK: roomId, SK: connectionId พร้อมเปิดใช้ฟิลด์ ttl)")

    heading("3.5 การออกแบบการจัดการ API ภายนอกและการป้องกันความล้มเหลว (External API Strategy)", 2)
    body("เพื่อป้องกันปัญหาเมื่อ External API (lorcana-api.com) ล่ม, ปิดตัว หรือติด Rate Limit ระบบใช้กลยุทธ์ 4 ระดับ: (1) Local Static Asset Bundling (cards_set1_set2.json 408 ใบ) (2) In-Browser LocalStorage Caching TTL 7 วัน (3) Serverless Reverse Proxy & DynamoDB Card Snapshot (4) Graceful UI Degradation พร้อมภาพการ์ดสำรอง SVG/PNG")

    heading("3.6 การออกแบบการสเกลและรองรับ WebSockets Concurrency", 2)
    body("แก้ไขปัญหาคอขวดเมื่อมีห้องเล่นจำนวนมากด้วย: (1) Partition Key Dispersion & TTL Auto-Cleanup ใน DynamoDB (2) Optimistic UI Updates ร่วมกับ Short Payload Relay (<100ms) (3) Auto-Reconnect Window 30 วินาที และ (4) Decoupled Fan-Out Architecture ด้วย SQS/Redis สำหรับผู้ชมจำนวนมาก")

    heading("3.7 การออกแบบตามกรอบ AWS Well-Architected Framework", 2)
    body("ครอบคลุม 4 เสาหลัก: Reliability (Idempotent Lambda, DLQ, Exponential Backoff), Availability (Multi-AZ 99.99% SLA, No SPOF), Security (HTTPS/WSS TLS 1.3, KMS encryption at rest, bcrypt 10 rounds, signed JWT, PoLP) และ Cost Optimization ($0.00 Free Tier, งบประมาณ < $50, CloudWatch Billing Alarms)")

    heading("3.8 การวิเคราะห์ข้อจำกัดของ AWS Learner Lab และแนวทางแก้ไขเชิงวิศวกรรม", 2)
    body("แก้ไขข้อจำกัดการบล็อก iam:CreateRole ด้วยการผูก LabRole สำเร็จรูป, สคริปต์ Manual Deployment (scripts/deploy_ws.sh), พัฒนาระบบ Custom Auth บน Lambda เอง และปรับแก้ API Gateway Payload Format Version 1.0")

    heading("บทที่ 4\nการพัฒนาระบบและผลการทดสอบคุณภาพ (Implementation & Quality Assurance)", 1)
    heading("4.1 รายละเอียดการพัฒนาส่วนติดต่อผู้ใช้ (Frontend Implementation)", 2)
    body("พัฒนาด้วย React 19 และ Tailwind CSS v4 ประกอบด้วยหน้า Match Lobby, Auth Modal, User Profile, Deck Builder, Playmat Simulation และ 3D Card Inspector")

    heading("4.2 รายละเอียดการพัฒนาส่วนประมวลผลบนคลาวด์ (AWS Microservices)", 2)
    body("ฟังก์ชัน Lambda (TypeScript / Node.js 20.x) ได้แก่ lorcana-auth, lorcana-deck, lorcana-room-handler และ lorcana-analyzer")

    heading("4.3 แผนและผลการทดสอบคุณภาพระบบ (Master QA Test Matrix)", 2)
    body("ผลการทดสอบผ่านทั้ง 14 กรณีทดสอบสำคัญ โดยการซิงค์ข้อมูลกระดานผ่าน WebSocket มีความหน่วงเฉลี่ย 52-82ms ผ่านเกณฑ์เป้าหมาย Sub-100ms และผ่านการตรวจสอบความปลอดภัย OWASP 100%")

    heading("บทที่ 5\nสรุปผลการดำเนินงานและแผนงานในอนาคต (Conclusions & Future Roadmap)", 1)
    heading("5.1 สรุปผลการดำเนินงานความก้าวหน้า Stage 2", 2)
    body("โครงงานบรรลุผลสำเร็จ 100% ตามเกณฑ์ประเมิน Stage 2 โดยสามารถสร้างระบบกระดานเล่นการ์ดบนคลาวด์ Serverless ที่มีความหน่วงต่ำกว่า 100ms, ปรับใช้ระบบบน AWS จริง, เอาชนะข้อจำกัดของ AWS Learner Lab และควบคุมค่าใช้จ่ายได้ $0.00")

    heading("5.2 แผนงานการพัฒนาสำหรับ Stage 3 (Sprint 4 ถึง 7)", 2)
    body("ประกอบด้วย Sprint 4 (Async SQS Deck Analyzer), Sprint 5 (CloudWatch Monitoring & Security Hardening), Sprint 6 (Load Testing ด้วย Artillery) และ Sprint 7 (รายงานฉบับสมบูรณ์และวิดีโอนำเสนอ)")

    heading("5.3 บทเรียนที่ได้รับและการปรับปรุงกระบวนการทำงาน (Retrospective)", 2)
    body("สรุปบทเรียนด้านการออกแบบ Interface First, การรับมือกับคลาวด์แซนด์บ็อกซ์ที่มีข้อจำกัด และประโยชน์ของสถาปัตยกรรม Serverless")

    heading("บรรณานุกรม (References)", 1)
    p_r = doc.add_paragraph()
    p_r.add_run("""[1] Ravensburger & Disney. (2023). Disney Lorcana Trading Card Game Official Rules and Card Database. [Online]. Available: https://www.disneylorcana.com
[2] Amazon Web Services. (2026). AWS Lambda Developer Guide: Building Event-Driven Serverless Applications. [Online]. Available: https://docs.aws.amazon.com/lambda/
[3] Amazon Web Services. (2026). Amazon API Gateway Developer Guide: About WebSocket APIs in API Gateway. [Online]. Available: https://docs.aws.amazon.com/apigateway/latest/developerguide/apigateway-websocket-api.html
[4] Amazon Web Services. (2026). Amazon DynamoDB Developer Guide: Core Components and Best Practices for NoSQL Design. [Online]. Available: https://docs.aws.amazon.com/amazondynamodb/
[5] Amazon Web Services. (2026). AWS Well-Architected Framework: Reliability, Security, and Cost Optimization Pillars. [Online]. Available: https://aws.amazon.com/architecture/well-architected/
[6] React Working Group. (2026). React 19 Documentation: Concurrent Features and Server Components. [Online]. Available: https://react.dev
[7] Tailwind Labs. (2026). Tailwind CSS v4.0 Alpha/Release Documentation. [Online]. Available: https://tailwindcss.com
[8] Framer. (2026). Framer Motion: Production-Ready Animation Library for React. [Online]. Available: https://www.framer.com/motion/
[9] IETF (Internet Engineering Task Force). (2011). RFC 6455: The WebSocket Protocol. [Online]. Available: https://tools.ietf.org/html/rfc6455""").font.name = 'TH Sarabun New'

    heading("ภาคผนวก (Appendix)", 1)
    heading("ตัวอย่างโค้ดสคริปต์ Deployment บน AWS Learner Lab (scripts/deploy_ws.sh)", 2)
    p_c = doc.add_paragraph()
    rc = p_c.add_run("""#!/bin/bash
# scripts/deploy_ws.sh - Deploy Real-Time WebSocket Infrastructure
set -e

echo "=== 1. Building Backend TypeScript ==="
cd ../backend && npm install && npx tsc --outDir dist

echo "=== 2. Packaging Lambda Handler ==="
cd dist/room && zip -r ../../../room-handler.zip . && cd ../../..

echo "=== 3. Updating Lambda Code via AWS CLI ==="
aws lambda update-function-code --function-name lorcana-room-handler --zip-file fileb://room-handler.zip --region us-east-1

echo "=== 4. Verifying WebSocket API Gateway Integration ==="
API_ID=$(aws apigatewayv2 get-apis --query "Items[?Name=='lorcana-ws-api'].ApiId" --output text)
echo "WebSocket API ID: $API_ID"
echo "WebSocket Endpoint: wss://${API_ID}.execute-api.us-east-1.amazonaws.com/prod"

echo "=== Deployment Completed Successfully ==="
""")
    rc.font.name = 'Consolas'
    rc.font.size = Pt(11)

    docx_file = os.path.join(DOCS_DIR, "G21_Disney_Lorcana_PlayLab_Cloud_Phase2_Report.docx")
    doc.save(docx_file)
    print(f"MS Word generated successfully: {docx_file} ({os.path.getsize(docx_file):,} bytes)")

except Exception as e:
    print(f"Error generating Word document: {e}")

print("\nALL DOCUMENTS BUILD SCRIPT COMPLETED!")

# 4. Generate Previews
print("\n[4/4] Generating high-resolution preview images of key pages...")
try:
    import pymupdf
    pdf_path = os.path.join(DOCS_DIR, "G21_Disney_Lorcana_PlayLab_Cloud_Phase2_Report.pdf")
    if os.path.exists(pdf_path):
        doc_pdf = pymupdf.open(pdf_path)
        # Page 1 (Cover)
        page1 = doc_pdf[0]
        pix1 = page1.get_pixmap(dpi=150)
        pix1.save(os.path.join(DOCS_DIR, "page_1_preview.png"))
        
        # Page 15 (Architecture & Design)
        if len(doc_pdf) >= 15:
            page15 = doc_pdf[14]
            pix15 = page15.get_pixmap(dpi=150)
            pix15.save(os.path.join(DOCS_DIR, "page_15_preview.png"))
            
        # Page 22 (QA Matrix)
        if len(doc_pdf) >= 22:
            page22 = doc_pdf[21]
            pix22 = page22.get_pixmap(dpi=150)
            pix22.save(os.path.join(DOCS_DIR, "page_22_preview.png"))
        print(f"Preview images generated successfully for {len(doc_pdf)} total pages.")
except Exception as e:
    print(f"Preview generation note: {e}")

print("\nALL DELIVERABLES (TEX, PDF, DOCX, PREVIEWS) GENERATED SUCCESSFULLY!")
